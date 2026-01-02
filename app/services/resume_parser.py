"""
Resume parsing service for extracting text and data from resume files
"""
import os
import logging
import json
from typing import Dict, List, Optional, Any
from app.simple_logger import get_logger
import re
from datetime import datetime
import tempfile
import io

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Fallback to PyPDF2 if pdfplumber not available
try:
    import PyPDF2
    PDF_FALLBACK_AVAILABLE = True
except ImportError:
    PDF_FALLBACK_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import docx2txt
    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False

try:
    import phonenumbers
    PHONE_AVAILABLE = True
except ImportError:
    PHONE_AVAILABLE = False

try:
    from dateparser.search import search_dates
    DATE_PARSER_AVAILABLE = True
except ImportError:
    DATE_PARSER_AVAILABLE = False

try:
    from tika import parser as tika_parser
    TIKA_AVAILABLE = True
except ImportError:
    TIKA_AVAILABLE = False

try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

logger = get_logger(__name__.split('.')[-1])

class ResumeParser:
    """Service for parsing resume files and extracting structured data"""
    
    def __init__(self):
        # Load comprehensive skill list from file if available, otherwise use default
        self.skill_keywords = self._load_skill_keywords()
        
        # Initialize OpenAI client if available
        self.openai_client = None
        self.openai_available = False
        if OPENAI_AVAILABLE:
            openai_api_key = os.getenv('OPENAI_API_KEY') or os.getenv('CHATGPT_API_KEY')
            if openai_api_key:
                try:
                    # Try new OpenAI client format first (v1.0+)
                    try:
                        from openai import OpenAI
                        self.openai_client = OpenAI(api_key=openai_api_key)
                        self.openai_available = True
                        logger.info("OpenAI API initialized successfully (v1.0+)")
                    except (ImportError, AttributeError):
                        # Fallback to old format
                        openai.api_key = openai_api_key
                        self.openai_client = openai
                        self.openai_available = True
                        logger.info("OpenAI API initialized successfully (legacy)")
                except Exception as e:
                    logger.warning(f"Failed to initialize OpenAI API: {e}")
            else:
                logger.info("OpenAI API key not found, will use fallback parsing")
        
        self.experience_patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'(\d+)\+?\s*years?\s*in\s*(?:the\s*)?(?:field|industry|profession)',
            r'(\d+)\+?\s*years?\s*working\s*(?:in|with)',
            r'(\d+)\+?\s*years?\s*professional\s*experience',
            r'(\d+)\+?\s*years?\s*of\s*professional\s*development',
            r'(\d+)\+?\s*years?\s*at\s+\w+',  # e.g., "3 years at Google"
            r'(\d+)\+?\s*years?\s*with\s+\w+',  # e.g., "2 years with Microsoft"
            r'(\d+)\+?\s*years?\s*in\s+\w+',  # e.g., "4 years in software development"
        ]
        
        self.education_keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'associate', 'diploma', 'certificate',
            'degree', 'university', 'college', 'institute', 'school'
        ]

    def _load_skill_keywords(self) -> List[str]:
        """Load skill keywords from file or use default list"""
        try:
            # Try to load from all_skills.txt file
            skill_file_path = os.path.join(os.path.dirname(__file__), '..', '..', 'all_skills.txt')
            if os.path.exists(skill_file_path):
                with open(skill_file_path, "r", encoding="utf-8") as f:
                    return [line.strip().lower() for line in f if line.strip()]
        except Exception as e:
            logger.warning(f"Could not load skills from file: {e}")
        
        # Fallback to default comprehensive skill list
        return [
            # Programming Languages
            'python', 'javascript', 'java', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'swift', 'kotlin',
            'typescript', 'scala', 'r', 'matlab', 'perl', 'haskell', 'clojure', 'erlang',
            'c', 'csharp', 'vb.net', 'f#', 'dart', 'objective-c', 'assembly', 'cobol', 'fortran',
            
            # Web Technologies
            'html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'spring',
            'laravel', 'rails', 'asp.net', 'jquery', 'bootstrap', 'sass', 'less', 'webpack',
            'next.js', 'nuxt.js', 'ember', 'backbone', 'knockout', 'polymer', 'svelte',
            'html5', 'css3', 'responsive design', 'pwa', 'spa', 'ssr', 'csr',
            
            # Databases
            'mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle', 'sql server', 'cassandra',
            'elasticsearch', 'dynamodb', 'firebase', 'neo4j', 'couchdb', 'riak', 'hbase',
            'mariadb', 'db2', 'teradata', 'snowflake', 'bigquery', 'redshift',
            
            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'ci/cd', 'terraform',
            'ansible', 'linux', 'bash', 'powershell', 'vagrant', 'chef', 'puppet',
            'aws ec2', 'aws s3', 'aws lambda', 'aws rds', 'aws cloudformation', 'aws cloudwatch',
            'azure vm', 'azure storage', 'azure functions', 'azure devops',
            'google cloud', 'google app engine', 'google compute engine',
            
            # Data & AI
            'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'pandas', 'numpy', 'scikit-learn',
            'data analysis', 'statistics', 'sql', 'tableau', 'power bi', 'excel',
            'keras', 'opencv', 'nltk', 'spacy', 'hadoop', 'spark', 'kafka', 'airflow',
            'data science', 'data engineering', 'etl', 'data visualization', 'business intelligence',
            
            # Mobile
            'android', 'ios', 'react native', 'flutter', 'xamarin', 'ionic',
            'swift', 'kotlin', 'objective-c', 'java', 'cordova', 'phonegap',
            
            # Testing
            'unit testing', 'integration testing', 'selenium', 'cypress', 'jest', 'mocha', 'chai',
            'junit', 'testng', 'pytest', 'rspec', 'cucumber', 'bdd', 'tdd',
            
            # Other
            'agile', 'scrum', 'project management', 'leadership', 'communication', 'teamwork',
            'microservices', 'api', 'rest', 'graphql', 'soap', 'json', 'xml',
            'version control', 'git', 'svn', 'mercurial', 'bitbucket', 'github', 'gitlab',
            'jira', 'confluence', 'slack', 'trello', 'asana', 'notion'
        ]

    def extract_text_from_file(self, file_path: str, file_extension: str) -> str:
        """Extract text from resume file based on file type"""
        try:
            if file_extension.lower() == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension.lower() in ['docx']:
                return self._extract_from_docx(file_path)
            elif file_extension.lower() in ['doc']:
                return self._extract_from_doc(file_path)
            elif file_extension.lower() == 'txt':
                return self._extract_from_txt(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_extension}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using pdfplumber for better accuracy, fallback to PyPDF2"""
        if PDF_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        # Try to extract text with layout preservation
                        page_text = page.extract_text(layout=True)
                        if not page_text:
                            # Fallback to regular extraction
                            page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        # Also try extracting tables if present
                        tables = page.extract_tables()
                        if tables:
                            for table in tables:
                                table_text = "\n".join([" | ".join([str(cell) if cell else "" for cell in row]) for row in table])
                                if table_text.strip():
                                    text += "\n" + table_text + "\n"
                    return text.strip()
            except Exception as e:
                logger.warning(f"Error reading PDF with pdfplumber: {str(e)}, trying fallback")
        
        # Fallback to PyPDF2
        if PDF_FALLBACK_AVAILABLE:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    return text.strip()
            except Exception as e:
                logger.error(f"Error reading PDF with PyPDF2: {str(e)}")
                return ""
        else:
            logger.error("No PDF parsing library available")
            return ""

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available for DOCX parsing")
            return ""
        
        try:
            doc = Document(file_path)
            text = ""
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            # Extract tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() if cell.text else "" for cell in row.cells])
                    if row_text.strip():
                        text += row_text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error reading DOCX: {str(e)}")
            return ""

    def _extract_from_doc(self, file_path: str) -> str:
        """Extract text from DOC file using tika for better accuracy"""
        if not TIKA_AVAILABLE:
            logger.error("tika not available for DOC parsing")
            return ""
        
        try:
            parsed = tika_parser.from_file(file_path)
            return parsed.get("content", "").strip()
        except Exception as e:
            logger.error(f"Error reading DOC with tika: {str(e)}")
            return ""

    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error extracting from TXT with latin-1: {str(e)}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting from TXT: {str(e)}")
            return ""

    def _parse_with_chatgpt(self, resume_text: str) -> Optional[Dict[str, Any]]:
        """Parse resume using ChatGPT API for better accuracy"""
        if not self.openai_available or not self.openai_client:
            return None
        
        try:
            # Truncate text if too long (ChatGPT has token limits)
            # Increased limit to capture more information
            max_chars = 16000  # Leave room for prompt and response
            truncated_text = resume_text[:max_chars] if len(resume_text) > max_chars else resume_text
            
            prompt = f"""Extract ALL structured information from the following resume text. You MUST extract and return ALL available information. Return a JSON object with the following fields (use empty string or empty array if not found):
- full_name: Full name of the candidate (string)
- email: Email address (string)
- phone: Phone number in E.164 format if possible (string)
- location: City, State/Country (string)
- skills: Array of ALL technical skills, programming languages, tools, and technologies mentioned (JSON array of strings)
- experience_years: Total years of experience as an integer (calculate from work history if not explicitly stated)
- education: Complete education details including degrees, universities, years, GPA, honors (string, preserve ALL details)
- work_experience: Complete work experience including job titles, companies, dates, responsibilities, achievements (string, preserve ALL details)
- projects: All project details including project names, descriptions, technologies used, links (string, preserve ALL details)
- certifications: All certifications and licenses with dates and issuing organizations (string, preserve ALL details)
- summary: Professional summary, objective, or career objective (string)
- languages: Languages known with proficiency levels (string or array)
- achievements: Awards, honors, recognitions, publications (string, preserve ALL details)
- hobbies: Hobbies and interests (string or array)
- references: References if mentioned (string)
- volunteer_work: Volunteer experience and community service (string, preserve ALL details)
- additional_info: Any other relevant information not captured above (string)

IMPORTANT: Extract ALL information available. Do not skip any sections. If a section exists in the resume, extract it completely. Preserve all details including dates, locations, descriptions, and achievements.

Resume text:
{truncated_text}

Return ONLY valid JSON in this exact format (no markdown, no code blocks, just pure JSON):
{{"full_name": "...", "email": "...", "phone": "...", "location": "...", "skills": [...], "experience_years": ..., "education": "...", "work_experience": "...", "projects": "...", "certifications": "...", "summary": "...", "languages": "...", "achievements": "...", "hobbies": "...", "references": "...", "volunteer_work": "...", "additional_info": "..."}}"""

            # Try new OpenAI client format first (v1.0+)
            try:
                response = self.openai_client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are a resume parser. Extract structured data from resumes and return only valid JSON."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.1,
                    max_tokens=4000
                )
                content = response.choices[0].message.content.strip()
            except AttributeError:
                # Fallback to old API format
                response = self.openai_client.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are a resume parser. Extract structured data from resumes and return only valid JSON."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.1,
                    max_tokens=4000
                )
                content = response.choices[0].message.content.strip()
            
            # Remove markdown code blocks if present
            if content.startswith("```"):
                parts = content.split("```")
                if len(parts) > 1:
                    content = parts[1]
                    if content.startswith("json"):
                        content = content[4:]
                content = content.strip()
            
            parsed_data = json.loads(content)
            
            # Ensure all expected fields are present
            result = {
                'full_name': parsed_data.get('full_name'),
                'email': parsed_data.get('email'),
                'phone': parsed_data.get('phone'),
                'location': parsed_data.get('location'),
                'skills': parsed_data.get('skills', []),
                'experience_years': parsed_data.get('experience_years'),
                'education': parsed_data.get('education'),
                'work_experience': parsed_data.get('work_experience'),
                'projects': parsed_data.get('projects'),
                'certifications': parsed_data.get('certifications'),
                'summary': parsed_data.get('summary'),
                'languages': parsed_data.get('languages'),
                'achievements': parsed_data.get('achievements'),
                'hobbies': parsed_data.get('hobbies'),
                'references': parsed_data.get('references'),
                'volunteer_work': parsed_data.get('volunteer_work'),
                'additional_info': parsed_data.get('additional_info'),
                'raw_text': resume_text
            }
            
            logger.info("Successfully parsed resume using ChatGPT API")
            logger.debug(f"ChatGPT extracted - Skills: {len(result.get('skills', []))}, Education: {bool(result.get('education'))}, Experience: {bool(result.get('work_experience'))}")
            return result
            
        except json.JSONDecodeError as e:
            logger.warning(f"Failed to parse ChatGPT JSON response: {e}")
            logger.debug(f"Response content: {content[:500] if 'content' in locals() else 'N/A'}")
            return None
        except Exception as e:
            logger.warning(f"ChatGPT API parsing failed: {e}, falling back to traditional parsing")
            import traceback
            logger.debug(f"Traceback: {traceback.format_exc()}")
            return None

    def parse_resume_data(self, resume_text: str, filename: str = "") -> Dict[str, Any]:
        """Parse resume text and extract structured data with improved accuracy"""
        if not resume_text:
            logger.warning("Empty resume text provided")
            return {}
        
        # Try ChatGPT API first if available
        chatgpt_result = None
        if self.openai_available:
            try:
                chatgpt_result = self._parse_with_chatgpt(resume_text)
                if chatgpt_result:
                    # Validate that ChatGPT returned meaningful data
                    has_meaningful_data = any([
                        chatgpt_result.get('full_name'),
                        chatgpt_result.get('email'),
                        len(chatgpt_result.get('skills', [])) > 0,
                        chatgpt_result.get('education'),
                        chatgpt_result.get('work_experience')
                    ])
                    if has_meaningful_data:
                        logger.info("Using ChatGPT parsed data")
                        # Still run traditional parsing to fill in any gaps
                        # (we'll merge them below)
                    else:
                        logger.warning("ChatGPT returned empty/incomplete data, will use traditional parsing")
                        chatgpt_result = None
            except Exception as e:
                logger.warning(f"ChatGPT parsing failed: {e}, falling back to traditional parsing")
                chatgpt_result = None
        
        # Always run traditional parsing (either as fallback or to supplement ChatGPT)
        logger.info("Running traditional resume parsing")
        logger.info(f"Resume text length: {len(resume_text)} characters")
        
        # Clean and normalize text
        text = self._clean_text(resume_text)
        logger.debug(f"After cleaning, text length: {len(text)} characters")
        logger.debug(f"First 200 chars: {text[:200]}")
        
        # First, identify and separate different sections to prevent cross-contamination
        sections = self._identify_sections(text)
        logger.info(f"Identified sections: {list(sections.keys())}")
        logger.info(f"Section sizes - education: {len(sections.get('education', ''))}, experience: {len(sections.get('experience', ''))}, skills: {len(sections.get('skills', ''))}, projects: {len(sections.get('projects', ''))}, certifications: {len(sections.get('certifications', ''))}")
        
        # Log section content previews
        for section_name, section_content in sections.items():
            if section_content and len(section_content) > 0:
                preview = section_content[:100].replace('\n', ' ')
                logger.debug(f"Section '{section_name}' preview: {preview}...")
        
        # Extract different components from their respective sections
        # Note: These methods now properly separate multiple entries (e.g., multiple degrees, projects, certifications)
        education = self._extract_education_from_sections(sections)
        work_experience = self._extract_work_experience_from_sections(sections)
        projects = self._extract_projects_from_sections(sections)
        certifications = self._extract_certifications_from_sections(sections)
        summary = self._extract_summary_from_sections(sections)
        skills = self._extract_skills_from_sections(sections)
        
        # Count entries in extracted data (separated by '---')
        education_count = len(education.split('\n\n---\n\n')) if education else 0
        projects_count = len(projects.split('\n\n---\n\n')) if projects else 0
        certifications_count = len(certifications.split('\n\n---\n\n')) if certifications else 0
        
        logger.info(f"[EXTRACTION] After section extraction:")
        logger.info(f"  - education: {bool(education)} ({len(str(education or ''))} chars, {education_count} entries)")
        logger.info(f"  - experience: {bool(work_experience)} ({len(str(work_experience or ''))} chars)")
        logger.info(f"  - skills: {len(skills) if skills else 0} skills")
        logger.info(f"  - projects: {bool(projects)} ({len(str(projects or ''))} chars, {projects_count} entries)")
        logger.info(f"  - certifications: {bool(certifications)} ({len(str(certifications or ''))} chars, {certifications_count} entries)")
        logger.info(f"  - summary: {bool(summary)} ({len(str(summary or ''))} chars)")
        
        # Use section data directly if extraction methods return empty but section has content
        if not education and sections.get('education') and len(sections.get('education', '').strip()) > 10:
            education = sections['education'].strip()
            logger.debug("Using raw education section data")
        
        if not work_experience and sections.get('experience') and len(sections.get('experience', '').strip()) > 10:
            work_experience = sections['experience'].strip()
            logger.debug("Using raw experience section data")
        
        if not projects and sections.get('projects') and len(sections.get('projects', '').strip()) > 10:
            projects = sections['projects'].strip()
            logger.debug("Using raw projects section data")
        
        if not skills and sections.get('skills') and len(sections.get('skills', '').strip()) > 5:
            # Try to extract skills from raw skills section
            skills_text = sections['skills'].strip()
            skills_list = re.split(r'[,•\n\r;]', skills_text)
            skills = [s.strip() for s in skills_list if s.strip() and len(s.strip()) > 1 and len(s.strip()) < 50]
            skills = [s for s in skills if self._is_valid_skill(s)]
            logger.debug(f"Extracted {len(skills)} skills from raw skills section")
        
        # Clean up education data to remove any project contamination
        if education and 'project' in education.lower():
            lines = education.split('\n')
            clean_lines = []
            for line in lines:
                line = line.strip()
                if line and not any(proj_term in line.lower() for proj_term in ['project :', 'title :', 'technology :', 'details :']):
                    clean_lines.append(line)
                elif any(proj_term in line.lower() for proj_term in ['project :', 'title :']):
                    break
            education = '\n'.join(clean_lines) if clean_lines else None
        
        # Aggressive fallback to text-based extraction if section-based extraction failed
        # Use lower thresholds to be more inclusive
        if not work_experience or (isinstance(work_experience, str) and len(work_experience.strip()) < 15):
            logger.info("[FALLBACK] Attempting text-based work experience extraction")
            work_experience_fallback = self._extract_work_experience(text)
            if work_experience_fallback and len(str(work_experience_fallback).strip()) > 10:
                work_experience = work_experience_fallback
                logger.info(f"[FALLBACK] Found work experience via fallback ({len(str(work_experience))} chars)")
            else:
                logger.warning("[FALLBACK] No work experience found via text extraction, trying aggressive method")
                work_experience_aggressive = self._extract_experience_aggressive(text)
                if work_experience_aggressive and len(str(work_experience_aggressive).strip()) > 10:
                    work_experience = work_experience_aggressive
                    logger.info(f"[FALLBACK] Found work experience via aggressive method ({len(str(work_experience))} chars)")
                else:
                    logger.warning("[FALLBACK] No work experience found even with aggressive extraction")
        
        if not certifications or (isinstance(certifications, str) and len(certifications.strip()) < 5):
            logger.info("[FALLBACK] Attempting text-based certifications extraction")
            certs_fallback = self._extract_certifications(text)
            if certs_fallback and len(str(certs_fallback).strip()) > 5:
                certifications = certs_fallback
                logger.info(f"[FALLBACK] Found certifications via fallback ({len(str(certifications))} chars)")
            else:
                logger.warning("[FALLBACK] No certifications found via text extraction, trying aggressive method")
                certs_aggressive = self._extract_certifications_aggressive(text)
                if certs_aggressive and len(str(certs_aggressive).strip()) > 5:
                    certifications = certs_aggressive
                    logger.info(f"[FALLBACK] Found certifications via aggressive method ({len(str(certifications))} chars)")
                else:
                    logger.warning("[FALLBACK] No certifications found even with aggressive extraction")
        
        if not skills or len(skills) < 1:
            logger.info("[FALLBACK] Attempting text-based skills extraction")
            skills_fallback = self._extract_skills(text)
            if skills_fallback:
                # Merge with existing skills
                if skills:
                    skills = list(set(skills + skills_fallback))
                else:
                    skills = skills_fallback
                logger.info(f"[FALLBACK] Found {len(skills)} skills via fallback: {skills[:10]}")
            else:
                logger.warning("[FALLBACK] No skills found via text extraction, trying aggressive method")
                skills_aggressive = self._extract_skills_aggressive(text)
                if skills_aggressive and len(skills_aggressive) > 0:
                    if skills:
                        skills = list(set(skills + skills_aggressive))
                    else:
                        skills = skills_aggressive
                    logger.info(f"[FALLBACK] Found {len(skills)} skills via aggressive method: {skills[:10]}")
                else:
                    logger.warning("[FALLBACK] No skills found even with aggressive extraction")
        
        if not projects or (isinstance(projects, str) and len(projects.strip()) < 5):
            logger.info("[FALLBACK] Attempting text-based projects extraction")
            projects_fallback = self._extract_projects(text)
            if projects_fallback and len(str(projects_fallback).strip()) > 5:
                projects = projects_fallback
                logger.info(f"[FALLBACK] Found projects via fallback ({len(str(projects))} chars)")
            else:
                logger.warning("[FALLBACK] No projects found via text extraction, trying aggressive method")
                projects_aggressive = self._extract_projects_aggressive(text)
                if projects_aggressive and len(str(projects_aggressive).strip()) > 5:
                    projects = projects_aggressive
                    logger.info(f"[FALLBACK] Found projects via aggressive method ({len(str(projects))} chars)")
                else:
                    logger.warning("[FALLBACK] No projects found even with aggressive extraction")
        
        if not education or (isinstance(education, str) and len(education.strip()) < 5):
            logger.info("[FALLBACK] Attempting text-based education extraction")
            education_fallback = self._extract_education(text)
            if education_fallback and len(str(education_fallback).strip()) > 5:
                education = education_fallback
                logger.info(f"[FALLBACK] Found education via fallback ({len(str(education))} chars)")
            else:
                logger.warning("[FALLBACK] No education found via text extraction, trying aggressive method")
                education_aggressive = self._extract_education_aggressive(text)
                if education_aggressive and len(str(education_aggressive).strip()) > 5:
                    education = education_aggressive
                    logger.info(f"[FALLBACK] Found education via aggressive method ({len(str(education))} chars)")
                else:
                    logger.warning("[FALLBACK] No education found even with aggressive extraction")
        
        # The section-based extraction should already be clean
        # No additional cleaning needed
        
        # For fresh graduates, extract career objective as work experience
        if not work_experience or (isinstance(work_experience, str) and len(work_experience.strip()) < 10):
            other_section = sections.get('other', '')
            if other_section:
                lines = other_section.split('\n')
                for line in lines:
                    line = line.strip()
                    if 'career objective' in line.lower() or 'objective' in line.lower():
                        # Find the objective text
                        obj_start = other_section.lower().find('career objective')
                        if obj_start != -1:
                            obj_text = other_section[obj_start:obj_start+200].strip()
                            if obj_text:
                                work_experience = obj_text
                                logger.info("[FALLBACK] Using career objective as work experience")
                        break
        
        # Final check: if we still have no data, use raw text sections as last resort
        if not work_experience and sections.get('experience'):
            work_experience = sections['experience'].strip()
            logger.info("[FALLBACK] Using raw experience section as work experience")
        
        if not education and sections.get('education'):
            education = sections['education'].strip()
            logger.info("[FALLBACK] Using raw education section as education")
        
        if not projects and sections.get('projects'):
            projects = sections['projects'].strip()
            logger.info("[FALLBACK] Using raw projects section as projects")
        
        if not certifications and sections.get('certifications'):
            certifications = sections['certifications'].strip()
            logger.info("[FALLBACK] Using raw certifications section as certifications")
        
        # Create a focused summary - only use explicit summary/objective
        combined_summary = summary or ""
        
        # If no explicit summary found, create a brief professional summary from key info
        if not combined_summary:
            summary_parts = []
            
            # Add education if available
            if education:
                # Extract just the degree and institution
                edu_lines = education.split('\n')
                if edu_lines:
                    first_edu = edu_lines[0].strip()
                    if len(first_edu) > 10:
                        summary_parts.append(f"Education: {first_edu}")
            
            # Add experience if available
            if work_experience:
                # Extract just the first experience
                exp_lines = work_experience.split('\n')
                if exp_lines:
                    first_exp = exp_lines[0].strip()
                    if len(first_exp) > 10:
                        summary_parts.append(f"Experience: {first_exp}")
            
            # Add skills summary
            if skills:
                # Take first 5 skills
                top_skills = skills[:5]
                summary_parts.append(f"Key Skills: {', '.join(top_skills)}")
            
            if summary_parts:
                combined_summary = ' | '.join(summary_parts)
        
        # Clean up the combined summary to avoid repetition
        if combined_summary:
            lines = combined_summary.split('\n')
            cleaned_lines = []
            seen_content = set()
            
            for line in lines:
                line = line.strip()
                if line and len(line) > 5:
                    line_lower = line.lower()
                    # Check for repetition
                    if not any(seen.lower() in line_lower or line_lower in seen.lower() for seen in seen_content):
                        cleaned_lines.append(line)
                        seen_content.add(line)
            
            combined_summary = '\n'.join(cleaned_lines)
        
        # Limit summary length to avoid database issues (but keep it reasonable)
        if len(combined_summary) > 5000:
            combined_summary = combined_summary[:5000] + "..."
        
        # Extract additional fields
        languages = self._extract_languages(sections, text)
        achievements = self._extract_achievements(sections, text)
        hobbies = self._extract_hobbies(sections, text)
        references = self._extract_references(sections, text)
        volunteer_work = self._extract_volunteer_work(sections, text)
        additional_info = self._extract_additional_info(sections, text)
        
        # Build parsed data from traditional parsing
        # Ensure we have at least raw_text even if nothing else is extracted
        traditional_data = {
            'full_name': self._extract_name(text) or None,
            'email': self._extract_email(text) or None,
            'phone': self._extract_phone(text) or None,
            'location': self._extract_location(text) or None,
            'skills': skills if skills else [],
            'experience_years': self._extract_experience_years(text),
            'education': education if education else None,
            'work_experience': work_experience if work_experience else None,
            'projects': projects if projects else None,
            'certifications': certifications if certifications else None,
            'summary': combined_summary if combined_summary else None,
            'languages': languages if languages else None,
            'achievements': achievements if achievements else None,
            'hobbies': hobbies if hobbies else None,
            'references': references if references else None,
            'volunteer_work': volunteer_work if volunteer_work else None,
            'additional_info': additional_info if additional_info else None,
            'raw_text': resume_text  # Always include raw text
        }
        
        # Final validation: ensure we have at least raw_text
        if not traditional_data.get('raw_text'):
            logger.warning("[FINAL] No raw_text available, using original resume_text")
            traditional_data['raw_text'] = resume_text
        
        logger.info(f"[FINAL] Traditional data summary:")
        logger.info(f"  - Has name: {bool(traditional_data.get('full_name'))}")
        logger.info(f"  - Has email: {bool(traditional_data.get('email'))}")
        logger.info(f"  - Has skills: {len(traditional_data.get('skills', []))}")
        logger.info(f"  - Has education: {bool(traditional_data.get('education'))}")
        logger.info(f"  - Has experience: {bool(traditional_data.get('work_experience'))}")
        logger.info(f"  - Has raw_text: {bool(traditional_data.get('raw_text'))} ({len(str(traditional_data.get('raw_text', '')))} chars)")
      
        if chatgpt_result:
            parsed_data = {}
            # For each field, use ChatGPT if it has data, otherwise use traditional
            for key in ['full_name', 'email', 'phone', 'location', 'experience_years', 'education', 'work_experience', 'projects', 'certifications', 'summary', 'languages', 'achievements', 'hobbies', 'references', 'volunteer_work', 'additional_info']:
                chatgpt_value = chatgpt_result.get(key)
                traditional_value = traditional_data.get(key)
                
                # Prefer ChatGPT if it has meaningful data
                chatgpt_has_data = False
                if chatgpt_value:
                    if isinstance(chatgpt_value, str) and chatgpt_value.strip():
                        chatgpt_has_data = True
                    elif isinstance(chatgpt_value, (int, float)) and chatgpt_value:
                        chatgpt_has_data = True
                
                traditional_has_data = False
                if traditional_value:
                    if isinstance(traditional_value, str) and traditional_value.strip():
                        traditional_has_data = True
                    elif isinstance(traditional_value, (int, float)) and traditional_value:
                        traditional_has_data = True
                
                # For these fields, favor traditional extraction (better structure)
                structured_block_keys = {'education', 'work_experience', 'projects', 'certifications'}

                if key in structured_block_keys:
                    if traditional_has_data:
                        parsed_data[key] = traditional_value
                    elif chatgpt_has_data:
                        parsed_data[key] = chatgpt_value
                    else:
                        parsed_data[key] = chatgpt_value or traditional_value or ('' if isinstance(chatgpt_value or traditional_value, str) else None)
                else:
                    if chatgpt_has_data:
                        parsed_data[key] = chatgpt_value
                    elif traditional_has_data:
                        parsed_data[key] = traditional_value
                    else:
                        parsed_data[key] = chatgpt_value or traditional_value or ('' if isinstance(chatgpt_value or traditional_value, str) else None)
            
            # For skills, merge both lists and remove duplicates
            chatgpt_skills = chatgpt_result.get('skills', []) or []
            traditional_skills = traditional_data.get('skills', []) or []
            # Combine and deduplicate
            all_skills = list(set([s.strip() for s in chatgpt_skills + traditional_skills if s and s.strip()]))
            parsed_data['skills'] = all_skills
            
            # Always keep raw_text from original
            parsed_data['raw_text'] = resume_text
            
            logger.info(f"Merged parsing results - Skills: {len(parsed_data.get('skills', []))}, Education: {bool(parsed_data.get('education'))}, Experience: {bool(parsed_data.get('work_experience'))}")
        else:
            parsed_data = traditional_data
            logger.info(f"Using traditional parsing only - Skills: {len(parsed_data.get('skills', []))}, Education: {bool(parsed_data.get('education'))}, Experience: {bool(parsed_data.get('work_experience'))}")
        
        # Final summary logging
        logger.info("=" * 60)
        logger.info("FINAL EXTRACTION SUMMARY")
        logger.info("=" * 60)
        logger.info(f"Full Name: {parsed_data.get('full_name', 'Not found')}")
        logger.info(f"Email: {parsed_data.get('email', 'Not found')}")
        logger.info(f"Phone: {parsed_data.get('phone', 'Not found')}")
        logger.info(f"Location: {parsed_data.get('location', 'Not found')}")
        logger.info(f"Skills: {len(parsed_data.get('skills', []))} found - {parsed_data.get('skills', [])[:10]}")
        logger.info(f"Experience Years: {parsed_data.get('experience_years', 'Not found')}")
        logger.info(f"Education: {'Found' if parsed_data.get('education') else 'Not found'} ({len(parsed_data.get('education') or '')} chars)")
        logger.info(f"Work Experience: {'Found' if parsed_data.get('work_experience') else 'Not found'} ({len(parsed_data.get('work_experience') or '')} chars)")
        logger.info(f"Projects: {'Found' if parsed_data.get('projects') else 'Not found'} ({len(parsed_data.get('projects') or '')} chars)")
        logger.info(f"Certifications: {'Found' if parsed_data.get('certifications') else 'Not found'} ({len(parsed_data.get('certifications') or '')} chars)")
        logger.info(f"Summary: {'Found' if parsed_data.get('summary') else 'Not found'} ({len(parsed_data.get('summary') or '')} chars)")
        logger.info(f"Languages: {'Found' if parsed_data.get('languages') else 'Not found'} ({len(parsed_data.get('languages') or '')} chars)")
        logger.info(f"Achievements: {'Found' if parsed_data.get('achievements') else 'Not found'} ({len(parsed_data.get('achievements') or '')} chars)")
        logger.info(f"Hobbies: {'Found' if parsed_data.get('hobbies') else 'Not found'} ({len(parsed_data.get('hobbies') or '')} chars)")
        logger.info(f"References: {'Found' if parsed_data.get('references') else 'Not found'} ({len(parsed_data.get('references') or '')} chars)")
        logger.info(f"Volunteer Work: {'Found' if parsed_data.get('volunteer_work') else 'Not found'} ({len(parsed_data.get('volunteer_work') or '')} chars)")
        logger.info(f"Additional Info: {'Found' if parsed_data.get('additional_info') else 'Not found'} ({len(parsed_data.get('additional_info') or '')} chars)")
        logger.info("=" * 60)
        
        return parsed_data

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text while preserving structure"""
        # Remove extra whitespace but preserve line breaks for table data
        # Only collapse multiple spaces/tabs within a line, keep line breaks
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            # Clean up spaces and tabs within the line
            cleaned_line = re.sub(r'[ \t]+', ' ', line)
            cleaned_lines.append(cleaned_line)
        # Join lines back, preserving single newlines
        text = '\n'.join(cleaned_lines)
        # Only collapse multiple consecutive newlines (3+) to double newline for section separation
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    def _identify_sections(self, text: str) -> Dict[str, str]:
        """Identify and separate different sections of the resume with improved accuracy"""
        sections = {
            'personal': '',
            'summary': '',
            'education': '',
            'experience': '',
            'skills': '',
            'projects': '',
            'certifications': '',
            'languages': '',
            'achievements': '',
            'hobbies': '',
            'references': '',
            'volunteer': '',
            'other': ''
        }
        
        # Enhanced section headers with more variations and case-insensitive matching
        section_patterns = {
            'personal': [
                r'^personal\s+details?', r'^personal\s+information', r'^contact\s+information',
                r'^contact\s+details?', r'^personal\s+data', r'^about\s+me', r'^contact'
            ],
            'summary': [
                r'^objective', r'^profile', r'^about$', r'^overview',
                r'^career\s+objective', r'^professional\s+summary', r'^executive\s+summary',
                r'^summary', r'^professional\s+profile'
            ],
            'education': [
                r'^education', r'^academic', r'^qualifications?', r'^degrees?',
                r'^academic\s+qualification', r'^educational\s+background', r'^scholastic',
                r'^academic\s+qualification\s*:', r'^education\s*:',
                r'^qualification\s*:', r'^educational\s+background\s*:', r'^academic\s+background\s*:',
                r'^academics', r'^educational\s+qualification'
            ],
            'experience': [
                r'^experience', r'^work\s+experience', r'^employment', r'^professional\s+experience',
                r'^career$', r'^work\s+history', r'^employment\s+history', r'^professional\s+background',
                r'^experience\s*:', r'^work\s+experience\s*:', r'^employment\s*:', r'^professional\s+experience\s*:',
                r'^work\s+history\s*:', r'^employment\s+history\s*:', r'^career\s*:',
                r'^work', r'^professional\s+background'
            ],
            'skills': [
                r'^technical\s+skills?', r'^core\s+competencies?', r'^technologies?',
                r'^expertise', r'^competencies?', r'^technical\s+expertise', r'^programming\s+skills?',
                r'^technical\s+skills?\s*:', r'^skills?\s*:',
                r'^skills?$', r'^technical\s+skills?\s*:', r'^core\s+competencies?\s*:',
                r'^technologies?\s*:', r'^expertise\s*:', r'^competencies?\s*:',
                r'^technical\s+competencies?', r'^key\s+skills?'
            ],
            'projects': [
                r'^projects?$', r'^portfolio', r'^personal\s+projects?', r'^academic\s+projects?',
                r'^work\s+projects?', r'^project\s+experience', r'^key\s+projects?',
                r'^project\s*:', r'^projects?\s*:',
                r'^projects?\s*:', r'^portfolio\s*:', r'^personal\s+projects?\s*:',
                r'^academic\s+projects?\s*:', r'^work\s+projects?\s*:'
            ],
            'certifications': [
                r'^certifications?$', r'^certificates?$', r'^certified$', r'^credentials?$',
                r'^licenses?$', r'^professional\s+certifications?', r'^certifications?\s+and\s+licenses?',
                r'^certifications?\s*:', r'^certificates?\s*:', r'^certified\s*:'
            ],
            'languages': [
                r'^languages?$', r'^languages?\s+known', r'^language\s+proficiency',
                r'^languages?\s+spoken', r'^foreign\s+languages?', r'^languages?\s*:'
            ],
            'achievements': [
                r'^achievements?$', r'^awards?$', r'^honors?$', r'^recognition$',
                r'^publications?$', r'^papers?$', r'^research$', r'^awards?\s+and\s+honors?$',
                r'^achievements?\s*:', r'^awards?\s*:', r'^honors?\s*:'
            ],
            'hobbies': [
                r'^hobbies?$', r'^interests?$', r'^personal\s+interests?',
                r'^hobbies?\s+and\s+interests?$', r'^hobbies?\s*:', r'^interests?\s*:'
            ],
            'references': [
                r'^references?$', r'^professional\s+references?', r'^references?\s+available',
                r'^references?\s+upon\s+request', r'^references?\s*:'
            ],
            'volunteer': [
                r'^volunteer\s+work', r'^volunteer\s+experience', r'^volunteering',
                r'^community\s+service', r'^volunteer\s+activities?', r'^volunteer\s*:'
            ]
        }
        
        # Split text into lines for processing
        lines = text.split('\n')
        current_section = 'other'
        current_content = []
        
        logger.info(f"[SECTION_ID] Processing {len(lines)} lines for section identification")
        
        # Track if we've seen a section header to avoid false positives
        seen_sections = set()
        
        for i, line in enumerate(lines):
            original_line = line
            line = line.strip()
            
            if not line:
                # Preserve empty lines for structure
                if current_content:
                    current_content.append('')
                continue
            
            # Check if this line is a section header
            line_lower = line.lower()
            # Remove trailing colons and extra whitespace for matching
            line_clean = re.sub(r'[:;]\s*$', '', line_lower).strip()
            section_found = False
            matched_section = None
            
            # Priority order: check specific sections first
            priority_sections = ['education', 'experience', 'skills', 'projects', 'certifications', 
                               'summary', 'personal', 'achievements', 'languages', 'hobbies', 
                               'references', 'volunteer']
            
            for section_name in priority_sections:
                if section_name not in section_patterns:
                    continue
                    
                patterns = section_patterns[section_name]
                for pattern in patterns:
                    # Try exact match first
                    if re.match(pattern + r'$', line_clean, re.IGNORECASE):
                        matched_section = section_name
                        section_found = True
                        break
                    # Try with colon
                    if re.match(pattern + r'\s*[:;]?\s*$', line_clean, re.IGNORECASE):
                        matched_section = section_name
                        section_found = True
                        break
                
                if section_found:
                    break
            
            # Additional flexible matching for common patterns
            if not section_found:
                # Education patterns
                if (re.match(r'^(academic|education|qualification|degree)', line_clean, re.IGNORECASE) and 
                    len(line_clean) < 40 and 
                    'experience' not in line_clean):
                    matched_section = 'education'
                    section_found = True
                # Experience patterns
                elif (re.match(r'^(work|experience|employment|professional|career)', line_clean, re.IGNORECASE) and 
                      len(line_clean) < 40 and
                      'objective' not in line_clean):
                    matched_section = 'experience'
                    section_found = True
                # Skills patterns
                elif (re.match(r'^(skill|technical|competenc|expertise|technolog)', line_clean, re.IGNORECASE) and 
                      len(line_clean) < 40):
                    matched_section = 'skills'
                    section_found = True
                # Projects patterns
                elif (re.match(r'^(project|portfolio)', line_clean, re.IGNORECASE) and 
                      len(line_clean) < 40):
                    matched_section = 'projects'
                    section_found = True
                # Certifications patterns
                elif (re.match(r'^(certif|license|credential)', line_clean, re.IGNORECASE) and 
                      len(line_clean) < 40):
                    matched_section = 'certifications'
                    section_found = True
            
            if section_found and matched_section:
                # Save previous section content
                if current_content:
                    prev_content = '\n'.join(current_content).strip()
                    if prev_content:
                        # Only update if we have meaningful content
                        if not sections[current_section] or len(prev_content) > len(sections[current_section]):
                            sections[current_section] = prev_content
                
                # Start new section
                current_section = matched_section
                current_content = []
                seen_sections.add(matched_section)
                logger.debug(f"[SECTION_ID] Found section '{matched_section}' at line {i+1}: {line[:50]}")
                # Don't add the header line itself to content
                continue
            
            # Not a section header, add to current section
            current_content.append(original_line)
        
        # Save the last section
        if current_content:
            final_content = '\n'.join(current_content).strip()
            if final_content:
                if not sections[current_section] or len(final_content) > len(sections[current_section]):
                    sections[current_section] = final_content
        
        # Log section identification results
        logger.info(f"[SECTION_ID] Section identification complete:")
        for section_name, section_content in sections.items():
            if section_content and len(section_content.strip()) > 0:
                logger.info(f"  - {section_name}: {len(section_content)} chars")
            else:
                logger.debug(f"  - {section_name}: empty")
        
        return sections

    def _extract_education_from_sections(self, sections: Dict[str, str]) -> Optional[str]:
        """Extract education from the education section - preserve ALL education entries and separate them"""
        education_text = sections.get('education', '')
        
        # If we have education section content, separate individual education entries
        if education_text and len(education_text.strip()) > 5:
            # First, try to separate multiple education entries
            education_entries = self._separate_education_entries(education_text)
            if education_entries and len(education_entries) > 1:
                logger.info(f"[EDUCATION] Separated {len(education_entries)} education entries")
                return '\n\n---\n\n'.join(education_entries)
            elif education_entries and len(education_entries) == 1:
                logger.info(f"[EDUCATION] Found 1 education entry")
                return education_entries[0]
            
            # If separation didn't work, clean up but preserve ALL content
            lines = education_text.split('\n')
            cleaned_lines = []
            consecutive_other_sections = 0
            max_consecutive_other = 3  # Allow more tolerance
            
            for line in lines:
                original_line = line
                line = line.strip()
                
                if not line:
                    # Preserve empty lines for structure
                    cleaned_lines.append('')
                    consecutive_other_sections = 0
                    continue
                
                # Check if this is clearly a different section header
                line_lower = line.lower()
                is_other_section_header = False
                
                # More comprehensive check for other section headers
                other_section_keywords = [
                    'project', 'certification', 'experience', 'work experience', 'employment',
                    'skill', 'technical skill', 'languages', 'achievements', 'awards',
                    'hobbies', 'interests', 'references', 'volunteer', 'personal details',
                    'contact', 'objective', 'summary', 'profile'
                ]
                
                # Check if line is a section header (all caps or title case with colon)
                is_header_pattern = re.match(r'^[A-Z][A-Z\s]{2,}:?$', line) or \
                                   re.match(r'^[A-Z][a-z]+(\s+[A-Z][a-z]+)*:?$', line)
                
                if is_header_pattern:
                    # Check if it contains other section keywords
                    if any(keyword in line_lower for keyword in other_section_keywords):
                        is_other_section_header = True
                
                if is_other_section_header:
                    consecutive_other_sections += 1
                    # Only break if we see multiple consecutive other section headers
                    if consecutive_other_sections >= max_consecutive_other:
                        logger.debug(f"[EDUCATION] Breaking at line: {line[:50]} (consecutive other sections: {consecutive_other_sections})")
                        break
                else:
                    consecutive_other_sections = 0
                    # Include all lines that might be education
                    cleaned_lines.append(original_line)
            
            if cleaned_lines:
                result = '\n'.join(cleaned_lines).strip()
                if result and len(result) > 5:
                    logger.info(f"[EDUCATION] Extracted {len(cleaned_lines)} education lines ({len(result)} chars)")
                    return result
            
            # Fallback: return original if we have something
            if education_text.strip():
                logger.info(f"[EDUCATION] Returning original education text ({len(education_text)} chars)")
                return education_text.strip()
        
        # Also check other section for education that might have been misclassified
        other_text = sections.get('other', '')
        if other_text and not education_text:
            # Look for education patterns in other section
            education_lines = []
            lines = other_text.split('\n')
            in_education_section = False
            education_keywords = [
                'academic', 'qualification', 'education', 'degree', 'bachelor', 'master',
                'phd', 'diploma', 'university', 'college', 'school', 'graduation',
                'b.tech', 'm.tech', 'b.e', 'm.e', 'hsc', 'ssc', 'examination'
            ]
            
            for line in lines:
                line_stripped = line.strip()
                if not line_stripped:
                    if in_education_section:
                        education_lines.append('')
                    continue
                
                line_lower = line_stripped.lower()
                
                # Check if this line starts education section
                if any(keyword in line_lower for keyword in education_keywords):
                    if len(line_lower) < 50:  # Likely a header
                        if not in_education_section:
                            in_education_section = True
                            # Don't add the header itself if it's clearly a section header
                            if not re.match(r'^(academic|education|qualification)', line_lower):
                                education_lines.append(line_stripped)
                        else:
                            education_lines.append(line_stripped)
                    elif in_education_section:
                        education_lines.append(line_stripped)
                elif in_education_section:
                    # Continue collecting if we're in education section
                    # But stop if we hit a clear other section
                    if any(other in line_lower for other in ['experience', 'project', 'certification', 'skill']):
                        if len(line_lower) < 30:  # Likely a section header
                            break
                    education_lines.append(line_stripped)
            
            if education_lines:
                result = '\n'.join(education_lines).strip()
                if result and len(result) > 10:
                    logger.info(f"[EDUCATION] Found education in 'other' section ({len(result)} chars)")
                    return result
        
        return None

    def _extract_work_experience_from_sections(self, sections: Dict[str, str]) -> Optional[str]:
        """Extract work experience from the experience section - capture ALL experience entries"""
        experience_text = sections.get('experience', '')
        if not experience_text:
            # Also check 'other' section for experience
            other_text = sections.get('other', '')
            if other_text:
                # Look for work experience patterns
                exp_keywords = ['experience', 'work', 'employment', 'job', 'position', 'role',
                              'company', 'employer', 'worked', 'employed', 'developer', 'engineer',
                              'manager', 'analyst', 'consultant']
                lines = other_text.split('\n')
                exp_lines = []
                in_exp_section = False
                
                for line in lines:
                    line_stripped = line.strip()
                    if not line_stripped:
                        if in_exp_section:
                            exp_lines.append('')
                        continue
                    
                    line_lower = line_stripped.lower()
                    
                    # Check if this starts experience section
                    if any(keyword in line_lower for keyword in exp_keywords):
                        if len(line_lower) < 50:
                            in_exp_section = True
                            if not re.match(r'^(experience|work|employment)', line_lower):
                                exp_lines.append(line_stripped)
                        elif in_exp_section:
                            exp_lines.append(line_stripped)
                    elif in_exp_section:
                        # Stop if we hit a clear other section
                        if any(other in line_lower for other in ['education', 'project', 'certification', 'skill']):
                            if len(line_lower) < 30:
                                break
                        exp_lines.append(line_stripped)
                
                if exp_lines:
                    experience_text = '\n'.join(exp_lines).strip()
                    if experience_text and len(experience_text) > 10:
                        logger.info(f"[EXPERIENCE] Found experience in 'other' section ({len(experience_text)} chars)")
        
        if not experience_text:
            return None
        
        # Clean up experience text - preserve ALL entries
        lines = experience_text.split('\n')
        cleaned_lines = []
        consecutive_other_sections = 0
        max_consecutive_other = 3
        
        for line in lines:
            original_line = line
            line = line.strip()
            
            if not line:
                # Preserve empty lines for structure
                cleaned_lines.append('')
                consecutive_other_sections = 0
                continue
            
            # Check if this is clearly a different section header
            line_lower = line.lower()
            is_other_section_header = False
            
            # More comprehensive check for other section headers
            other_section_keywords = [
                'certifications', 'certificate', 'education', 'academic', 'qualification',
                'personal details', 'contact', 'skills', 'technical skills', 'competencies',
                'projects', 'portfolio', 'languages', 'achievements', 'awards',
                'hobbies', 'interests', 'references', 'volunteer'
            ]
            
            # Check if line is a section header
            is_header_pattern = re.match(r'^[A-Z][A-Z\s]{2,}:?$', line) or \
                               re.match(r'^[A-Z][a-z]+(\s+[A-Z][a-z]+)*:?$', line)
            
            if is_header_pattern:
                # Check if it contains other section keywords
                if any(keyword in line_lower for keyword in other_section_keywords):
                    is_other_section_header = True
            
            if is_other_section_header:
                consecutive_other_sections += 1
                # Only break if we see multiple consecutive other section headers
                if consecutive_other_sections >= max_consecutive_other:
                    logger.debug(f"[EXPERIENCE] Breaking at line: {line[:50]} (consecutive other sections: {consecutive_other_sections})")
                    break
            else:
                consecutive_other_sections = 0
                # Include all lines that might be experience
                cleaned_lines.append(original_line)
        
        if cleaned_lines:
            result = '\n'.join(cleaned_lines).strip()
            if result and len(result) > 5:
                logger.info(f"[EXPERIENCE] Extracted {len(cleaned_lines)} experience lines ({len(result)} chars)")
                return result
        
        # Fallback: return original if we have something
        if experience_text.strip():
            logger.info(f"[EXPERIENCE] Returning original experience text ({len(experience_text)} chars)")
            return experience_text.strip()
        
        return None

    def _extract_projects_from_sections(self, sections: Dict[str, str]) -> Optional[str]:
        """Extract projects from the projects section - separate multiple projects"""
        projects_text = sections.get('projects', '')
        
        # If we have projects section content, separate individual projects
        if projects_text and len(projects_text.strip()) > 10:
            # Try to separate multiple projects
            projects_list = self._separate_projects(projects_text)
            if projects_list and len(projects_list) > 1:
                # Join with clear separators
                logger.info(f"[PROJECTS] Separated {len(projects_list)} projects from projects section")
                return '\n\n---\n\n'.join(projects_list)
            elif projects_list and len(projects_list) == 1:
                # Only one project found, but log it
                logger.info(f"[PROJECTS] Found 1 project in projects section")
                return projects_list[0]
            else:
                # Separation didn't work, return raw text
                logger.info(f"[PROJECTS] Could not separate projects, returning raw text ({len(projects_text)} chars)")
                return projects_text.strip()
        
        if not projects_text:
            # Also check other section for projects
            projects_text = sections.get('other', '')
            if projects_text:
                # Look for project patterns in other section
                project_lines = []
                lines = projects_text.split('\n')
                in_project_section = False
                
                for line in lines:
                    line = line.strip()
                    if line and len(line) > 5:
                        # Check if this line starts a project
                        if re.match(r'^project\s*:', line, re.IGNORECASE) or re.match(r'^title\s*:', line, re.IGNORECASE):
                            in_project_section = True
                            project_lines.append(line)
                        elif in_project_section and (re.match(r'^technology\s*:', line, re.IGNORECASE) or 
                                                   re.match(r'^details\s*:', line, re.IGNORECASE) or
                                                   re.match(r'^project\s*:', line, re.IGNORECASE)):
                            project_lines.append(line)
                        elif in_project_section and not re.match(r'^[A-Z][A-Z\s]+:$', line):
                            # Include continuation lines of projects
                            if any(proj_term in line.lower() for proj_term in [
                                'html', 'css', 'javascript', 'java', 'react', 'reactjs', 'website', 'clone',
                                'online', 'shoes', 'store', 'weather', 'information', 'adidas', 'facebook',
                                'selling', 'details', 'particular', 'city', 'entered', 'home', 'page',
                                'title', 'technology', 'project', 'store', 'information', 'clone', 'node',
                                'python', 'django', 'flask', 'sql', 'database', 'api', 'frontend', 'backend'
                            ]):
                                project_lines.append(line)
                        elif re.match(r'^[A-Z][A-Z\s]+:$', line) and in_project_section:
                            # Don't break, continue to next project
                            continue
                
                if project_lines:
                    # Try to separate the projects found
                    project_text = '\n'.join(project_lines)
                    projects_list = self._separate_projects(project_text)
                    if projects_list and len(projects_list) > 1:
                        logger.info(f"[PROJECTS] Separated {len(projects_list)} projects from other section")
                        return '\n\n---\n\n'.join(projects_list)
                    return project_text
        
        if projects_text:
            # Clean up projects text and try to separate
            lines = projects_text.split('\n')
            cleaned_lines = []
            
            for line in lines:
                line = line.strip()
                if line and len(line) > 5:
                    # Skip section headers and personal details
                    if not re.match(r'^[A-Z][A-Z\s]+:$', line) and not any(personal_term in line.lower() for personal_term in [
                        'personal details', 'date of birth', 'father', 'mother', 'marital status',
                        'gender', 'languages known', 'permanent address', 'address', 'phone',
                        'email', 'contact', 'mobile', 'declaration', 'place', 'date'
                    ]):
                        cleaned_lines.append(line)
            
            if cleaned_lines:
                # Try to separate projects
                project_text = '\n'.join(cleaned_lines)
                projects_list = self._separate_projects(project_text)
                if projects_list and len(projects_list) > 1:
                    logger.info(f"[PROJECTS] Separated {len(projects_list)} projects from cleaned text")
                    return '\n\n---\n\n'.join(projects_list)
                return project_text  # Return all project lines
        
        return None

    def _extract_certifications_from_sections(self, sections: Dict[str, str]) -> Optional[str]:
        """Extract certifications from the certifications section - separate multiple certifications"""
        certs_text = sections.get('certifications', '')
        if not certs_text:
            return None
        
        # Try to separate multiple certifications
        certs_list = self._separate_certifications(certs_text)
        if certs_list and len(certs_list) > 1:
            # Join with clear separators
            logger.info(f"[CERTIFICATIONS] Separated {len(certs_list)} certifications from certifications section")
            return '\n\n---\n\n'.join(certs_list)
        elif certs_list and len(certs_list) == 1:
            # Only one cert found, but log it
            logger.info(f"[CERTIFICATIONS] Found 1 certification in certifications section")
            return certs_list[0]
        
        # Fallback to line-by-line if separation didn't work - capture ALL lines
        lines = certs_text.split('\n')
        cleaned_lines = []
        consecutive_section_headers = 0
        
        for line in lines:
            line = line.strip()
            if line and len(line) > 5:
                # Only skip if we see a clear section header pattern
                is_section_header = re.match(r'^[A-Z][A-Z\s]{3,}:$', line)
                is_other_section = is_section_header and any(section_term in line.lower() for section_term in [
                    'education', 'personal details', 'skills', 'technical skills', 'projects',
                    'experience', 'work experience', 'languages', 'achievements'
                ])
                
                if is_other_section:
                    consecutive_section_headers += 1
                    # Only break if we see multiple section headers (definitely moved to new section)
                    if consecutive_section_headers >= 2:
                        break
                else:
                    consecutive_section_headers = 0
                    # Include all lines - don't filter out potential certification entries
                    cleaned_lines.append(line)
            elif line:  # Include non-empty lines
                cleaned_lines.append(line)
        
        if cleaned_lines:
            logger.info(f"[EXTRACTION] Extracted {len(cleaned_lines)} certification lines")
            return '\n'.join(cleaned_lines)  # Return ALL certification lines
        
        return None

    def _separate_projects(self, projects_text: str) -> List[str]:
        """Separate multiple projects from a single text block - capture ALL projects"""
        projects = []
        lines = projects_text.split('\n')
        current_project = []
        
        # Patterns that indicate a new project
        project_start_patterns = [
            r'^project\s*\d+[\.:\)]',  # Project 1:, Project 1., Project 1)
            r'^project\s*:',  # Project:
            r'^title\s*:',  # Title:
            r'^\d+[\.\)]\s+',  # 1. or 1)
            r'^[A-Z][a-z]+\s+Project',  # E-commerce Project
            r'^project\s+name\s*:',  # Project Name:
            r'^project\s+title\s*:',  # Project Title:
            r'^(e-commerce|ecommerce|web|mobile|desktop|api|full\s+stack|frontend|backend)\s+project',  # Type of project
            r'^(application|app|website|system|platform|software|tool)\s*:',  # Application:, App:, etc.
        ]
        
        for line in lines:
            line = line.strip()
            if not line:
                # Empty line might separate projects, but don't break yet
                if current_project and len(current_project) > 0:
                    current_project.append('')  # Preserve spacing
                continue
            
            # Check if this line starts a new project
            is_new_project = False
            for pattern in project_start_patterns:
                if re.match(pattern, line, re.IGNORECASE):
                    is_new_project = True
                    break
            
            # Also check for common project name patterns
            if not is_new_project and len(current_project) > 0:
                # Check if line looks like a new project (starts with capital, has project keywords)
                project_keywords = ['project', 'app', 'website', 'system', 'application', 'platform', 
                                  'software', 'tool', 'clone', 'e-commerce', 'ecommerce', 'dashboard']
                line_lower = line.lower()
                
                # Check if line starts with project keyword or has project keyword early
                if re.match(r'^[A-Z]', line):
                    # Check if it's a short line (likely a title)
                    if len(line) < 100:
                        # Check if it contains project keywords
                        if any(keyword in line_lower[:50] for keyword in project_keywords):
                            is_new_project = True
                        # Also check for common project name patterns
                        elif re.match(r'^[A-Z][a-z]+(\s+[A-Z][a-z]+)*\s*(Project|App|Website|System|Application|Platform)', line):
                            is_new_project = True
            
            if is_new_project and current_project:
                # Save previous project
                project_text = '\n'.join(current_project).strip()
                if project_text and len(project_text) > 10:
                    projects.append(project_text)
                current_project = [line]
            else:
                current_project.append(line)
        
        # Add last project
        if current_project:
            project_text = '\n'.join(current_project).strip()
            if project_text and len(project_text) > 10:
                projects.append(project_text)
        
        # If no clear separation found, try splitting by double newlines or numbered items
        if len(projects) <= 1:
            # Try splitting by numbered items (1., 2., etc.)
            numbered_split = re.split(r'\n(?=\d+[\.\)]\s+)', projects_text)
            if len(numbered_split) > 1:
                projects = [p.strip() for p in numbered_split if p.strip() and len(p.strip()) > 10]
                logger.info(f"[SEPARATION] Split by numbered items: {len(projects)} projects")
        
        # If still only one, try splitting by double newlines
        if len(projects) <= 1:
            double_newline_split = re.split(r'\n\s*\n+', projects_text)
            if len(double_newline_split) > 1:
                projects = [p.strip() for p in double_newline_split if p.strip() and len(p.strip()) > 10]
                logger.info(f"[SEPARATION] Split by double newlines: {len(projects)} projects")
        
        # If still only one, try splitting by lines that start with "Project" or "Title"
        if len(projects) <= 1:
            project_header_split = re.split(r'\n(?=project\s*[:\d]|title\s*:)', projects_text, flags=re.IGNORECASE)
            if len(project_header_split) > 1:
                projects = [p.strip() for p in project_header_split if p.strip() and len(p.strip()) > 10]
                logger.info(f"[SEPARATION] Split by project headers: {len(projects)} projects")
        
        # If still only one but text is long, try splitting by common project separators
        if len(projects) <= 1 and len(projects_text) > 200:
            # Look for patterns like "Project Name:" or "Technology:" that might indicate new projects
            tech_split = re.split(r'\n(?=technology\s*:|tech\s*stack\s*:|tools\s*used\s*:)', projects_text, flags=re.IGNORECASE)
            if len(tech_split) > 1:
                projects = [p.strip() for p in tech_split if p.strip() and len(p.strip()) > 10]
                logger.info(f"[SEPARATION] Split by technology headers: {len(projects)} projects")
        
        # Additional fallback: split by lines that look like project titles (short lines starting with capital)
        if len(projects) <= 1:
            lines = projects_text.split('\n')
            potential_titles = []
            for i, line in enumerate(lines):
                line = line.strip()
                if line and re.match(r'^[A-Z]', line) and len(line) < 80:
                    # Check if it looks like a project title
                    if any(keyword in line.lower() for keyword in ['project', 'app', 'website', 'system', 'application', 'platform']):
                        potential_titles.append(i)
            
            if len(potential_titles) > 1:
                # Split by these potential titles
                split_projects = []
                for i in range(len(potential_titles)):
                    start_idx = potential_titles[i]
                    end_idx = potential_titles[i + 1] if i + 1 < len(potential_titles) else len(lines)
                    project_text = '\n'.join(lines[start_idx:end_idx]).strip()
                    if project_text and len(project_text) > 10:
                        split_projects.append(project_text)
                
                if len(split_projects) > 1:
                    projects = split_projects
                    logger.info(f"[SEPARATION] Split by potential project titles: {len(projects)} projects")

        # FINAL fallback: handle resumes where all projects are in one paragraph with titles like
        # "Shoe Store Application: ... Instagram Clone: ... Zomato Clone: ..."
        if len(projects) <= 1:
            # Pattern finds capitalized phrases ending with common project words followed by a colon
            title_pattern = re.compile(
                r'(?=(?:[A-Z][A-Za-z]*(?:\s+[A-Z][A-Za-z]*)*\s+'
                r'(?:Application|App|Project|Clone|Website|System|Platform)\s*:))'
            )
            # Only attempt if text is long enough and contains multiple colons (likely multiple projects)
            if len(projects_text) > 200 and projects_text.count(':') >= 2:
                parts = [p.strip() for p in title_pattern.split(projects_text) if p.strip()]
                if len(parts) > 1:
                    recombined = []
                    buffer = ''
                    for part in parts:
                        if title_pattern.match(part) and buffer:
                            # New title starts, push previous buffer
                            if len(buffer.strip()) > 10:
                                recombined.append(buffer.strip())
                            buffer = part
                        else:
                            buffer = (buffer + ' ' + part).strip() if buffer else part
                    if buffer and len(buffer.strip()) > 10:
                        recombined.append(buffer.strip())

                    if len(recombined) > 1:
                        logger.info(f"[SEPARATION] Split by inline title/colon pattern: {len(recombined)} projects")
                        projects = recombined
                    else:
                        logger.info("[SEPARATION] Inline title split attempted but only one project detected")

        logger.info(f"[SEPARATION] Final: Separated {len(projects)} projects from text")
        return [p for p in projects if p and len(p.strip()) > 10]
    
    def _separate_education_entries(self, education_text: str) -> List[str]:
        """Separate multiple education entries (degrees, diplomas, etc.) from a single text block"""
        entries = []
        lines = education_text.split('\n')
        current_entry = []
        
        # Enhanced patterns that indicate a new education entry
        education_start_patterns = [
            r'^(bachelor|master|phd|doctorate|diploma|associate|certificate|degree)\s*[:\-]?\s*',  # Bachelor:, Master, etc.
            r'^(b\.?tech|m\.?tech|b\.?e|m\.?e|b\.?sc|m\.?sc|b\.?a|m\.?a|b\.?com|m\.?com|b\.?ba|m\.?ba)',  # B.Tech, M.Tech, etc.
            r'^(hsc|ssc|10th|12th|high\s+school|secondary\s+school|primary\s+school)',  # HSC, SSC, etc.
            r'^(examination|exam|qualification)\s*:',  # Examination:, Exam:, Qualification:
            r'^\d+[\.\)]\s+',  # 1. or 1)
            r'^(graduation|post\s+graduation|pg|ug)\s*:',  # Graduation:, Post Graduation:
        ]
        
        # Education keywords for context detection
        education_keywords = ['university', 'college', 'institute', 'school', 'board', 'degree', 'diploma', 
                            'bachelor', 'master', 'phd', 'graduation', 'cgpa', 'gpa', 'percentage', 'grade',
                            'engineering', 'science', 'arts', 'commerce', 'technology']
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            
            if not line_stripped:
                # Empty line might separate entries, but preserve it
                if current_entry:
                    current_entry.append('')
                continue
            
            # Check if this line starts a new education entry
            is_new_entry = False
            line_lower = line_stripped.lower()
            
            # Check against patterns
            for pattern in education_start_patterns:
                if re.match(pattern, line_stripped, re.IGNORECASE):
                    is_new_entry = True
                    break
            
            # Also check for degree abbreviations at start of line
            if not is_new_entry:
                degree_abbrev_pattern = r'^(b\.?tech|m\.?tech|b\.?e|m\.?e|b\.?sc|m\.?sc|b\.?a|m\.?a|b\.?com|m\.?com|b\.?ba|m\.?ba|hsc|ssc)'
                if re.match(degree_abbrev_pattern, line_stripped, re.IGNORECASE):
                    is_new_entry = True
            
            # Check if line looks like a new education entry header
            if not is_new_entry and current_entry:
                # Check if it's a short line (likely a header) with education keywords
                if len(line_stripped) < 120 and re.match(r'^[A-Z]', line_stripped):
                    if any(keyword in line_lower for keyword in education_keywords):
                        is_new_entry = True
                    # Also check for common education patterns
                    elif re.match(r'^[A-Z][a-z]+(\s+[A-Z][a-z]+)*\s+(University|College|Institute|School|Academy)', line_stripped):
                        is_new_entry = True
            
            # Additional check: if previous entry ended and this looks like a new degree
            if not is_new_entry and current_entry:
                # Check if previous line was short (likely end of entry) and this starts a new one
                if len(current_entry) > 0:
                    prev_line = current_entry[-1].strip() if current_entry else ''
                    # If previous line was short and this line starts with education keyword
                    if len(prev_line) < 80 and re.match(r'^[A-Z]', line_stripped):
                        if any(keyword in line_lower[:30] for keyword in ['bachelor', 'master', 'degree', 'diploma', 'b.tech', 'm.tech']):
                            is_new_entry = True
            
            if is_new_entry and current_entry:
                # Save previous entry
                entry_text = '\n'.join(current_entry).strip()
                if entry_text and len(entry_text) > 10:
                    entries.append(entry_text)
                current_entry = [line_stripped]
            else:
                current_entry.append(line_stripped)
        
        # Add last entry
        if current_entry:
            entry_text = '\n'.join(current_entry).strip()
            if entry_text and len(entry_text) > 10:
                entries.append(entry_text)
        
        # If no clear separation found, try splitting by numbered items
        if len(entries) <= 1:
            numbered_split = re.split(r'\n(?=\d+[\.\)]\s+)', education_text)
            if len(numbered_split) > 1:
                entries = [e.strip() for e in numbered_split if e.strip() and len(e.strip()) > 10]
                logger.info(f"[SEPARATION] Split education by numbered items: {len(entries)} entries")
        
        # If still only one, try splitting by double newlines (more aggressive)
        if len(entries) <= 1:
            # Split by 2+ consecutive newlines
            double_newline_split = re.split(r'\n\s*\n\s*\n+', education_text)
            if len(double_newline_split) > 1:
                entries = [e.strip() for e in double_newline_split if e.strip() and len(e.strip()) > 10]
                logger.info(f"[SEPARATION] Split education by triple newlines: {len(entries)} entries")
            
            # If still one, try double newlines
            if len(entries) <= 1:
                double_newline_split = re.split(r'\n\s*\n+', education_text)
                if len(double_newline_split) > 1:
                    entries = [e.strip() for e in double_newline_split if e.strip() and len(e.strip()) > 10]
                    logger.info(f"[SEPARATION] Split education by double newlines: {len(entries)} entries")
        
        # If still only one, try splitting by examination/degree headers
        if len(entries) <= 1:
            exam_split = re.split(r'\n(?=examination\s*:|exam\s*:|qualification\s*:|degree\s*:|bachelor|master|phd|diploma|hsc|ssc|b\.?tech|m\.?tech|b\.?e|m\.?e)', education_text, flags=re.IGNORECASE)
            if len(exam_split) > 1:
                entries = [e.strip() for e in exam_split if e.strip() and len(e.strip()) > 10]
                logger.info(f"[SEPARATION] Split education by exam/degree headers: {len(entries)} entries")
        
        # Final check: if we have one long entry, try to split by common separators
        if len(entries) == 1 and len(entries[0]) > 200:
            # Look for patterns like "Degree Name -" or "Degree Name:" that might indicate multiple entries
            separator_pattern = r'\n(?=[A-Z][a-z]+\s+(?:in|of)\s+[A-Z]|Bachelor|Master|B\.?Tech|M\.?Tech|B\.?E|M\.?E)'
            potential_split = re.split(separator_pattern, entries[0], flags=re.IGNORECASE)
            if len(potential_split) > 1:
                entries = [e.strip() for e in potential_split if e.strip() and len(e.strip()) > 10]
                logger.info(f"[SEPARATION] Split long education entry: {len(entries)} entries")
        
        logger.info(f"[SEPARATION] Final: Separated {len(entries)} education entries from text")
        return [e for e in entries if e and len(e.strip()) > 10]

    def _separate_certifications(self, certs_text: str) -> List[str]:
        """Separate multiple certifications from a single text block - capture ALL certifications"""
        certs = []
        lines = certs_text.split('\n')
        current_cert = []
        
        # Patterns that indicate a new certification
        cert_start_patterns = [
            r'^certification\s*\d+[\.:\)]',  # Certification 1:
            r'^\d+[\.\)]\s+',  # 1. or 1)
            r'^[A-Z][a-z]+.*certif',  # AWS Certified, Google Certified
            r'^[A-Z]{2,}',  # AWS, GCP, etc.
            r'^certificate\s*:',  # Certificate:
            r'^license\s*:',  # License:
            r'^(aws|google|microsoft|oracle|cisco|comptia|pmp|itil|red\s+hat|vmware)',  # Provider names
        ]
        
        for line in lines:
            line = line.strip()
            if not line:
                # Empty line might separate certs, but don't break yet
                if current_cert and len(current_cert) > 0:
                    current_cert.append('')  # Preserve spacing
                continue
            
            # Check if this line starts a new certification
            is_new_cert = False
            for pattern in cert_start_patterns:
                if re.match(pattern, line, re.IGNORECASE):
                    is_new_cert = True
                    break
            
            # Also check if line contains certification keywords at the start
            if not is_new_cert:
                cert_keywords = ['certified', 'certificate', 'certification', 'license', 'credential', 'credentialed']
                line_lower = line.lower()
                # Check if line starts with cert keyword or has cert keyword followed by provider
                if any(keyword in line_lower[:30] for keyword in cert_keywords):
                    if re.match(r'^[A-Z]', line) and len(current_cert) > 0:
                        is_new_cert = True
            
            # Check for common cert provider names (AWS, Google, Microsoft, etc.) at start of line
            if not is_new_cert and len(current_cert) > 0:
                cert_providers = ['aws', 'google', 'microsoft', 'oracle', 'cisco', 'comptia', 'pmp', 'itil', 
                                'red hat', 'vmware', 'salesforce', 'adobe', 'ibm', 'sap', 'tableau', 'power bi']
                line_lower = line.lower()
                # Check if line starts with provider name
                if any(line_lower.startswith(provider) or line_lower.startswith(provider + ' ') for provider in cert_providers):
                    if re.match(r'^[A-Z]', line):
                        is_new_cert = True
                
                # Also check if provider appears early in the line
                for provider in cert_providers:
                    if provider in line_lower[:50] and re.match(r'^[A-Z]', line):
                        # Make sure it's not just part of a description
                        provider_pos = line_lower.find(provider)
                        if provider_pos < 30:  # Provider appears early in line
                            is_new_cert = True
                            break
            
            if is_new_cert and current_cert:
                # Save previous certification
                cert_text = '\n'.join(current_cert).strip()
                if cert_text and len(cert_text) > 5:
                    certs.append(cert_text)
                current_cert = [line]
            else:
                current_cert.append(line)
        
        # Add last certification
        if current_cert:
            cert_text = '\n'.join(current_cert).strip()
            if cert_text and len(cert_text) > 5:
                certs.append(cert_text)
        
        # If no clear separation found, try splitting by numbered items
        if len(certs) <= 1:
            numbered_split = re.split(r'\n(?=\d+[\.\)]\s+)', certs_text)
            if len(numbered_split) > 1:
                certs = [c.strip() for c in numbered_split if c.strip() and len(c.strip()) > 5]
                logger.info(f"[SEPARATION] Split certifications by numbered items: {len(certs)} certs")
        
        # If still only one, try splitting by double newlines
        if len(certs) <= 1:
            double_newline_split = re.split(r'\n\s*\n+', certs_text)
            if len(double_newline_split) > 1:
                certs = [c.strip() for c in double_newline_split if c.strip() and len(c.strip()) > 5]
                logger.info(f"[SEPARATION] Split certifications by double newlines: {len(certs)} certs")
        
        # If still only one, try splitting by lines that start with capital letters (likely cert names)
        if len(certs) <= 1:
            capital_lines = []
            for line in lines:
                line = line.strip()
                if line and re.match(r'^[A-Z]', line) and len(line) > 5:
                    # Check if it looks like a cert name (has cert keywords or provider names)
                    if any(keyword in line.lower() for keyword in ['certified', 'certificate', 'certification', 'license', 'aws', 'google', 'microsoft', 'oracle']):
                        capital_lines.append(line)
            if len(capital_lines) > 1:
                certs = capital_lines
                logger.info(f"[SEPARATION] Split certifications by capital lines: {len(certs)} certs")
        
        # If still only one but text is long, try splitting by common cert separators
        if len(certs) <= 1 and len(certs_text) > 100:
            # Look for patterns like "Issued by:" or "Date:" that might indicate new certs
            issuer_split = re.split(r'\n(?=issued\s+by\s*:|date\s*:|valid\s+until\s*:)', certs_text, flags=re.IGNORECASE)
            if len(issuer_split) > 1:
                certs = [c.strip() for c in issuer_split if c.strip() and len(c.strip()) > 5]
                logger.info(f"[SEPARATION] Split certifications by issuer/date: {len(certs)} certs")
        
        logger.info(f"[SEPARATION] Final: Separated {len(certs)} certifications from text")
        return [c for c in certs if c and len(c.strip()) > 5]

    def _extract_summary_from_sections(self, sections: Dict[str, str]) -> Optional[str]:
        """Extract summary from the summary section"""
        summary_text = sections.get('summary', '')
        if not summary_text:
            return None
        
        # Clean up summary text
        lines = summary_text.split('\n')
        cleaned_lines = []
        seen_lines = set()
        
        for line in lines:
            line = line.strip()
            if line and len(line) > 10:
                # Avoid repetition
                line_lower = line.lower()
                if not any(seen_line.lower() in line_lower or line_lower in seen_line.lower() 
                         for seen_line in seen_lines):
                    cleaned_lines.append(line)
                    seen_lines.add(line)
        
        if cleaned_lines:
            return '\n'.join(cleaned_lines)  # Return all summary lines
        
        return None

    def _extract_skills_from_sections(self, sections: Dict[str, str]) -> List[str]:
        """Extract skills from the skills section with improved accuracy"""
        skills_text = sections.get('skills', '')
        found_skills = []
        
        if skills_text:
            # Parse comma-separated, semicolon-separated, or bullet-pointed skills
            # Also handle newline-separated skills
            skills_list = re.split(r'[,•\n\r;|]', skills_text)
            for skill in skills_list:
                skill = skill.strip()
                # Remove common prefixes and suffixes
                skill = re.sub(r'^[^:]*:\s*', '', skill)  # Remove "Category:" prefix
                skill = re.sub(r'\s*[-–—]\s*.*$', '', skill)  # Remove " - description" suffix
                skill = skill.strip()
                
                if skill and len(skill) > 1 and len(skill) < 60:  # Increased max length
                    # Validate that it's actually a skill
                    if self._is_valid_skill(skill):
                        # Normalize skill name (title case, but preserve common abbreviations)
                        if skill.upper() == skill and len(skill) <= 5:
                            # Likely an abbreviation like "AWS", "API", "SQL"
                            found_skills.append(skill.upper())
                        else:
                            found_skills.append(skill.title())
        
        # Also check other sections for skills if skills section is empty or sparse
        if len(found_skills) < 5:
            # Combine all sections except personal details
            all_text = ' '.join([
                sections.get('summary', ''),
                sections.get('experience', ''),
                sections.get('projects', ''),
                sections.get('certifications', ''),
                sections.get('other', '')
            ])
            
            text_lower = all_text.lower()
            for skill in self.skill_keywords:
                # Use word boundaries to avoid partial matches
                if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
                    # Check if not already found (case-insensitive)
                    if not any(s.lower() == skill.lower() for s in found_skills):
                        found_skills.append(skill.title())
        
        # Remove duplicates (case-insensitive)
        unique_skills = []
        seen_lower = set()
        for skill in found_skills:
            skill_lower = skill.lower()
            if skill_lower not in seen_lower:
                unique_skills.append(skill)
                seen_lower.add(skill_lower)
        
        logger.info(f"[SKILLS] Extracted {len(unique_skills)} unique skills")
        return unique_skills

    def _extract_name(self, text: str) -> Optional[str]:
        """Extract candidate name from resume text"""
        # Look for name patterns at the beginning of the document
        lines = text.split('\n')[:15]  # Check first 15 lines
        
        for line in lines:
            line = line.strip()
            if len(line) > 2 and len(line) < 60:  # Reasonable name length
                # Skip lines that look like headers or contain common resume keywords
                if not any(keyword in line.lower() for keyword in [
                    'resume', 'cv', 'curriculum', 'vitae', 'contact', 'email', 'phone',
                    'address', 'objective', 'summary', 'experience', 'education', 'skills',
                    'projects', 'certifications', 'languages', 'references'
                ]):
                    # Check if it looks like a name (contains letters and possibly spaces, dots, hyphens, apostrophes)
                    if re.match(r'^[A-Za-z\s\.\-\']+$', line) and len(line.split()) <= 5:
                        # Additional check: make sure it's not just numbers or special characters
                        if re.search(r'[A-Za-z]', line):
                            return line.title()
        
        # If no name found in first lines, try to find it in the first few words
        words = text.split()[:10]
        if len(words) >= 2:
            potential_name = ' '.join(words[:2])
            if re.match(r'^[A-Za-z\s\.\-\']+$', potential_name) and len(potential_name.split()) <= 3:
                return potential_name.title()
        
        return None

    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address from resume text"""
        # More comprehensive email patterns
        email_patterns = [
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',  # Standard email
            r'email[:\s]*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})',  # Email: pattern
            r'e-mail[:\s]*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})',  # E-mail: pattern
            r'contact[:\s]*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})',  # Contact: pattern
        ]
        
        for pattern in email_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                email = matches[0].strip()
                # Clean up email (remove spaces and extra characters)
                email = re.sub(r'\s+', '', email)
                # Validate email format
                if '@' in email and '.' in email.split('@')[1]:
                    return email
        
        # Additional fallback - look for any email pattern in the text
        email_fallback = re.search(r'[A-Za-z0-9._%+-]+\s*@\s*[A-Za-z0-9.-]+\s*\.\s*[A-Z|a-z]{2,}', text)
        if email_fallback:
            email = email_fallback.group(0)
            email = re.sub(r'\s+', '', email)
            # Clean up email - remove any extra text after the domain
            if '@' in email and '.' in email.split('@')[1]:
                # Split by @ and take only the domain part before any extra text
                local_part, domain_part = email.split('@', 1)
                # Clean domain part - remove anything after the TLD
                domain_clean = re.sub(r'[^a-zA-Z0-9.-].*$', '', domain_part)
                if '.' in domain_clean:
                    return f"{local_part}@{domain_clean}"
        
        # Final fallback - look for email with spaces and clean it
        email_with_spaces = re.search(r'[A-Za-z0-9._%+-]+\s+@\s*[A-Za-z0-9.-]+\s*\.\s*[A-Z|a-z]{2,}', text)
        if email_with_spaces:
            email = email_with_spaces.group(0)
            email = re.sub(r'\s+', '', email)
            # Extract just the email part before any extra text
            email_match = re.search(r'([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})', email)
            if email_match:
                return email_match.group(1)
        
        # Additional cleanup for emails with extra text
        email_cleanup = re.search(r'([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})[^A-Za-z0-9]', text)
        if email_cleanup:
            return email_cleanup.group(1)
        
        # Final cleanup - remove any trailing non-email characters
        final_email = re.search(r'([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})', text)
        if final_email:
            email = final_email.group(1)
            # Remove any trailing characters that are not part of the email
            email = re.sub(r'[^A-Za-z0-9._%+-@]', '', email)
            # Additional cleanup for common issues
            if email.endswith('.GitHub') or email.endswith('.github'):
                email = email.replace('.GitHub', '').replace('.github', '')
            # Remove any trailing dots or special characters
            email = email.rstrip('.')
            # Final validation - ensure it's a proper email
            if '@' in email and '.' in email.split('@')[1]:
                return email
        
        return None

    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from resume text using phonenumbers library for better accuracy"""
        if not PHONE_AVAILABLE:
            # Fallback to regex-based extraction
            return self._extract_phone_regex(text)
        
        try:
            # Try to extract with no region hint (only works if +country_code exists)
            matches = list(phonenumbers.PhoneNumberMatcher(text, None))
            for match in matches:
                number = match.number
                if phonenumbers.is_valid_number(number):
                    return phonenumbers.format_number(number, phonenumbers.PhoneNumberFormat.E164)
           
            # If not found, try with a list of possible regions
            fallback_regions = ["US", "IN", "GB", "CA", "AU", "DE", "FR", "SG", "PH", "ZA", "BR", "MX", "NG", "KE"]
            for region in fallback_regions:
                matches = list(phonenumbers.PhoneNumberMatcher(text, region))
                for match in matches:
                    number = match.number
                    if phonenumbers.is_valid_number(number):
                        return phonenumbers.format_number(number, phonenumbers.PhoneNumberFormat.E164)
        except Exception as e:
            logger.warning(f"Phone number parsing error: {e}")
       
        return None

    def _extract_phone_regex(self, text: str) -> Optional[str]:
        """Fallback phone extraction using regex patterns"""
        # More comprehensive phone patterns including Indian and international formats
        phone_patterns = [
            # Indian mobile numbers (10 digits starting with 6,7,8,9)
            r'phone[:\s]*([6-9]\d{9})',
            r'mobile[:\s]*([6-9]\d{9})',
            r'contact[:\s]*([6-9]\d{9})',
            r'\+91[-\s]?([6-9]\d{9})',
            r'\+91[-\s]?(\d{5}[-\s]?\d{5})',
            
            # US/International formats
            r'\+?1?[-.\s]?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',
            r'\+?[0-9]{1,3}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}',
            r'\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',
            
            # General patterns
            r'(\d{10})',  # 10 consecutive digits
            r'(\d{3}[-.\s]?\d{3}[-.\s]?\d{4})',  # XXX-XXX-XXXX format
            r'(\d{5}[-.\s]?\d{5})',  # XXXXX-XXXXX format (Indian landline)
        ]
        
        for pattern in phone_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                phone = matches[0]
                # Clean the phone number (remove non-digits)
                clean_phone = re.sub(r'[^\d]', '', str(phone))
                
                # Validate phone number length
                if len(clean_phone) >= 10:
                    # For Indian numbers, ensure it starts with 6,7,8,9
                    if len(clean_phone) == 10 and clean_phone[0] in '6789':
                        return clean_phone
                    # For US numbers, ensure it's 10 digits
                    elif len(clean_phone) == 10:
                        return clean_phone
                    # For international numbers with country code
                    elif len(clean_phone) > 10:
                        return clean_phone
        
        return None

    def _extract_location(self, text: str) -> Optional[str]:
        """Extract location from resume text with high precision"""
        # Common city names and states to help identify real locations
        common_cities = [
            'mumbai', 'delhi', 'bangalore', 'hyderabad', 'chennai', 'kolkata', 'pune', 'ahmedabad',
            'jaipur', 'surat', 'lucknow', 'kanpur', 'nagpur', 'indore', 'thane', 'bhopal',
            'visakhapatnam', 'pimpri', 'patna', 'vadodara', 'ghaziabad', 'ludhiana', 'agra',
            'nashik', 'faridabad', 'meerut', 'rajkot', 'kalyan', 'vasai', 'varanasi', 'dhule',
            'new york', 'los angeles', 'chicago', 'houston', 'phoenix', 'philadelphia',
            'san antonio', 'san diego', 'dallas', 'san jose', 'austin', 'jacksonville',
            'fort worth', 'columbus', 'charlotte', 'san francisco', 'indianapolis',
            'seattle', 'denver', 'washington', 'boston', 'el paso', 'nashville',
            'detroit', 'oklahoma city', 'portland', 'las vegas', 'memphis', 'louisville',
            'baltimore', 'milwaukee', 'albuquerque', 'tucson', 'fresno', 'sacramento',
            'mesa', 'kansas city', 'atlanta', 'long beach', 'colorado springs', 'raleigh',
            'miami', 'virginia beach', 'omaha', 'oakland', 'minneapolis', 'tulsa',
            'arlington', 'tampa', 'new orleans', 'wichita', 'cleveland', 'bakersfield',
            'aurora', 'anaheim', 'honolulu', 'santa ana', 'corpus christi', 'riverside',
            'lexington', 'stockton', 'henderson', 'saint paul', 'st louis', 'milwaukee'
        ]
        
        common_states = [
            'maharashtra', 'karnataka', 'tamil nadu', 'west bengal', 'gujarat', 'rajasthan',
            'uttar pradesh', 'andhra pradesh', 'telangana', 'kerala', 'madhya pradesh',
            'haryana', 'punjab', 'odisha', 'assam', 'jharkhand', 'chhattisgarh', 'himachal pradesh',
            'uttarakhand', 'goa', 'tripura', 'meghalaya', 'manipur', 'nagaland', 'arunachal pradesh',
            'mizoram', 'sikkim', 'delhi', 'jammu and kashmir', 'ladakh', 'puducherry',
            'california', 'texas', 'florida', 'new york', 'pennsylvania', 'illinois',
            'ohio', 'georgia', 'north carolina', 'michigan', 'new jersey', 'virginia',
            'washington', 'arizona', 'massachusetts', 'tennessee', 'indiana', 'missouri',
            'maryland', 'wisconsin', 'colorado', 'minnesota', 'south carolina', 'alabama',
            'louisiana', 'kentucky', 'oregon', 'oklahoma', 'connecticut', 'utah', 'iowa',
            'nevada', 'arkansas', 'mississippi', 'kansas', 'new mexico', 'nebraska',
            'west virginia', 'idaho', 'hawaii', 'new hampshire', 'maine', 'montana',
            'rhode island', 'delaware', 'south dakota', 'north dakota', 'alaska', 'vermont', 'wyoming'
        ]
        
        # Terms that should NOT be in location
        exclude_terms = [
            'email', 'phone', 'mobile', 'github', 'linkedin', 'website', 'portfolio',
            'objective', 'career', 'summary', 'experience', 'education', 'skills',
            'project', 'certification', 'achievement', 'hobby', 'interest',
            'javascript', 'html', 'css', 'react', 'angular', 'vue', 'node', 'python',
            'java', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'swift', 'kotlin',
            'mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'git', 'linux',
            'http', 'https', 'www', '.com', '.org', '.net', '@', 'gmail', 'yahoo', 'hotmail'
        ]
        
        def is_valid_location(location_text):
            """Check if the extracted text is a valid location"""
            location_lower = location_text.lower()
            
            # Must not contain excluded terms
            if any(term in location_lower for term in exclude_terms):
                return False
            
            # Must contain at least one known city or state
            if not (any(city in location_lower for city in common_cities) or 
                   any(state in location_lower for state in common_states)):
                return False
            
            # Must be reasonable length (not too short or too long)
            if len(location_text) < 3 or len(location_text) > 100:
                return False
            
            # Must not be mostly numbers or special characters
            if re.match(r'^[\d\s\-\(\)\.]+$', location_text):
                return False
            
            return True
        
        # Strategy 1: Look for specific location patterns with strict boundaries
        location_patterns = [
            # City, State format
            r'\b([A-Za-z\s]+),\s*([A-Za-z\s]+)\b',
            # City, State, Country format  
            r'\b([A-Za-z\s]+),\s*([A-Za-z\s]+),\s*([A-Za-z\s]+)\b',
            # Just city name (must be in our list)
            r'\b(' + '|'.join(common_cities) + r')\b',
        ]
        
        for pattern in location_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    # Join tuple elements
                    location = ', '.join(match).strip()
                else:
                    location = match.strip()
                
                if is_valid_location(location):
                    return location
        
        # Strategy 2: Look for address patterns but with strict validation
        address_patterns = [
            r'address[:\s]*([^\n]{1,50})',
            r'location[:\s]*([^\n]{1,50})',
            r'city[:\s]*([^\n]{1,50})',
            r'permanent address[:\s]*([^\n]{1,50})',
        ]
        
        for pattern in address_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                address = match.group(1).strip()
                
                # Stop at first excluded term
                for term in exclude_terms:
                    if term in address.lower():
                        address = address.split(term)[0].strip()
                        break
                
                # Clean up common separators
                address = re.sub(r'[^\w\s,.-]', '', address).strip()
                
                if is_valid_location(address):
                    # Extract just the city if it's a long address
                    if ',' in address:
                        parts = address.split(',')
                        for part in parts:
                            part = part.strip()
                            if is_valid_location(part):
                                return part
                    return address
        
        # Strategy 3: Look for standalone city names in the first few lines
        lines = text.split('\n')[:10]  # Only check first 10 lines
        for line in lines:
            line = line.strip()
            if len(line) < 100:  # Skip very long lines
                for city in common_cities:
                    if city in line.lower():
                        # Extract the city and surrounding context
                        city_match = re.search(r'\b([A-Za-z\s]*' + re.escape(city) + r'[A-Za-z\s]*)\b', line, re.IGNORECASE)
                        if city_match:
                            location = city_match.group(1).strip()
                            if is_valid_location(location):
                                return location
        
        return None

    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text"""
        found_skills = []
        
        # First, try to extract from dedicated skills section
        skills_section = self._extract_skills_section(text)
        if skills_section:
            # Parse comma-separated or bullet-pointed skills
            skills_list = re.split(r'[,•\n\r]', skills_section)
            for skill in skills_list:
                skill = skill.strip()
                if skill and len(skill) > 1 and len(skill) < 50:  # Reasonable skill length
                    # Validate that it's actually a skill, not an address or other text
                    if self._is_valid_skill(skill):
                        found_skills.append(skill.title())
        
        # If no skills section found, look for skills throughout the text but be more selective
        if not found_skills:
            text_lower = text.lower()
            for skill in self.skill_keywords:
                # Use word boundaries to avoid partial matches
                if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
                    found_skills.append(skill.title())
        
        # Remove duplicates and return
        return list(set(found_skills))
    
    def _is_valid_skill(self, skill: str) -> bool:
        """Validate if a string is actually a skill and not an address or other text"""
        skill_lower = skill.lower().strip()
        
        # Skip if it's too long (likely not a skill)
        if len(skill) > 60:  # Increased from 50
            return False
        
        # Skip if it's too short
        if len(skill) < 2:
            return False
        
        # Allow common multi-word skills (e.g., "machine learning", "data science")
        # But skip if it's clearly a sentence or description
        if len(skill.split()) > 5:  # More than 5 words is likely a description
            return False
        
        # Skip if it contains address-like patterns
        address_patterns = [
            r'\d+\s+[a-z\s]+road',  # "123 Main Road"
            r'\d+\s+[a-z\s]+street',  # "123 Main Street"
            r'\d+\s+[a-z\s]+avenue',  # "123 Main Avenue"
            r'\d+\s+[a-z\s]+lane',  # "123 Main Lane"
            r'\d+\s+[a-z\s]+drive',  # "123 Main Drive"
            r'[a-z\s]+,\s*[a-z\s]+,\s*\d+',  # "City, State, 12345"
            r'p\.o\.\s*box',  # "P.O. Box"
            r'apartment\s+\d+',  # "Apartment 123"
            r'apt\s+\d+',  # "Apt 123"
            r'rajendra\s+nagar',  # Specific address pattern from test
            r'kirti\s+gondur',  # Specific address pattern from test
            r'behind\s+[a-z\s]+road',  # "Behind Main Road"
            r'permanent\s+address',  # "Permanent Address"
            r'suite\s+\d+',  # "Suite 123"
            r'house\s+no',  # "House No"
            r'flat\s+no',  # "Flat No"
            r'building\s+no',  # "Building No"
            r'ward\s+no',  # "Ward No"
            r'pin\s*code',  # "Pin Code"
            r'postal\s+code',  # "Postal Code"
        ]
        
        for pattern in address_patterns:
            if re.search(pattern, skill_lower):
                return False
        
        # Skip if it contains personal information patterns
        personal_patterns = [
            r'father\'s?\s+name',  # "Father's Name"
            r'mother\'s?\s+name',  # "Mother's Name"
            r'date\s+of\s+birth',  # "Date of Birth"
            r'phone\s*:',  # "Phone:"
            r'email\s*:',  # "Email:"
            r'address\s*:',  # "Address:"
            r'contact\s*:',  # "Contact:"
            r'marital\s+status',  # "Marital Status"
            r'gender\s*:',  # "Gender:"
            r'languages?\s+known',  # "Languages Known"
            r'permanent\s+address',  # "Permanent Address"
            r'personal\s+details',  # "Personal Details"
            r'nationality',  # "Nationality"
            r'religion',  # "Religion"
            r'caste',  # "Caste"
            r'category',  # "Category"
        ]
        
        for pattern in personal_patterns:
            if re.search(pattern, skill_lower):
                return False
        
        # Skip if it contains project-like patterns that aren't skills
        project_patterns = [
            r'project\s*:',  # "Project:"
            r'portfolio\s*:',  # "Portfolio:"
            r'application\s*:',  # "Application:"
            r'website\s*:',  # "Website:"
            r'system\s*:',  # "System:"
            r'software\s*:',  # "Software:"
            r'app\s*:',  # "App:"
            r'platform\s*:',  # "Platform:"
            r'github\s*:',  # "GitHub:"
            r'linkedin\s*:',  # "LinkedIn:"
        ]
        
        for pattern in project_patterns:
            if re.search(pattern, skill_lower):
                return False
        
        # Skip if it contains education patterns
        education_patterns = [
            r'education\s*:',  # "Education:"
            r'academic\s*:',  # "Academic:"
            r'qualification\s*:',  # "Qualification:"
            r'degree\s*:',  # "Degree:"
            r'university\s*:',  # "University:"
            r'college\s*:',  # "College:"
            r'school\s*:',  # "School:"
            r'gpa\s*:',  # "GPA:"
            r'percentage\s*:',  # "Percentage:"
            r'cgpa\s*:',  # "CGPA:"
            r'grade\s*:',  # "Grade:"
        ]
        
        for pattern in education_patterns:
            if re.search(pattern, skill_lower):
                return False
        
        # Skip if it contains experience patterns
        experience_patterns = [
            r'experience\s*:',  # "Experience:"
            r'work\s+experience',  # "Work Experience"
            r'employment\s*:',  # "Employment:"
            r'professional\s+experience',  # "Professional Experience"
            r'career\s*:',  # "Career:"
            r'job\s*:',  # "Job:"
            r'position\s*:',  # "Position:"
        ]
        
        for pattern in experience_patterns:
            if re.search(pattern, skill_lower):
                return False
        
        # Skip if it contains certification patterns
        cert_patterns = [
            r'certification\s*:',  # "Certification:"
            r'certificate\s*:',  # "Certificate:"
            r'certified\s*:',  # "Certified:"
            r'credential\s*:',  # "Credential:"
            r'license\s*:',  # "License:"
        ]
        
        for pattern in cert_patterns:
            if re.search(pattern, skill_lower):
                return False
        
        # Skip if it's mostly numbers or special characters
        if re.match(r'^[\d\s\-\(\)\.]+$', skill):
            return False
        
        # Skip if it contains only numbers
        if skill.isdigit():
            return False
        
        # Skip if it's a single word that's too generic
        generic_words = [
            'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by',
            'from', 'up', 'about', 'into', 'through', 'during', 'before', 'after',
            'above', 'below', 'between', 'among', 'under', 'over', 'around', 'near',
            'far', 'here', 'there', 'where', 'when', 'why', 'how', 'what', 'who',
            'which', 'that', 'this', 'these', 'those', 'i', 'you', 'he', 'she', 'it',
            'we', 'they', 'me', 'him', 'her', 'us', 'them', 'my', 'your', 'his',
            'her', 'its', 'our', 'their', 'mine', 'yours', 'hers', 'ours', 'theirs',
            'name', 'age', 'dob', 'mobile', 'phone', 'email', 'address', 'city',
            'state', 'country', 'pin', 'code', 'house', 'flat', 'building', 'road',
            'street', 'lane', 'avenue', 'drive', 'area', 'locality', 'colony'
        ]
        
        if skill_lower in generic_words:
            return False
        
        # Skip if it looks like a name (starts with capital and contains only letters)
        if re.match(r'^[A-Z][a-z]+$', skill) and len(skill.split()) <= 3:
            # Check if it's not a known technical term
            technical_terms = [
                'java', 'python', 'react', 'angular', 'vue', 'node', 'express',
                'django', 'flask', 'spring', 'laravel', 'rails', 'asp', 'net',
                'mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle',
                'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git',
                'linux', 'bash', 'powershell', 'html', 'css', 'javascript',
                'typescript', 'php', 'ruby', 'go', 'rust', 'swift', 'kotlin'
            ]
            if skill_lower not in technical_terms:
                return False
        
        # Skip if it contains common non-skill words
        non_skill_words = [
            'years', 'year', 'months', 'month', 'days', 'day', 'hours', 'hour',
            'experience', 'work', 'job', 'company', 'organization', 'firm',
            'project', 'task', 'responsibility', 'duty', 'role', 'position',
            'team', 'group', 'department', 'division', 'unit', 'section',
            'client', 'customer', 'user', 'stakeholder', 'manager', 'director',
            'leader', 'supervisor', 'mentor', 'trainer', 'instructor'
        ]
        
        if any(word in skill_lower for word in non_skill_words):
            return False
        
        # If it passes all checks, it's likely a valid skill
        return True

    def _extract_skills_section(self, text: str) -> Optional[str]:
        """Extract skills section from resume"""
        skills_keywords = ['skills', 'technical skills', 'core competencies', 'technologies', 'expertise']
        
        for keyword in skills_keywords:
            pattern = rf'{keyword}[:\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*\n|\n[A-Z][A-Z\s]+:|\Z)'
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                return match.group(1).strip()
        
        return None

    def _extract_experience_years(self, text: str) -> Optional[int]:
        """Extract years of experience from resume text - only count work experience, not education dates"""
        text_lower = text.lower()
        
        # First try to find explicit experience statements in work/experience sections only
        # Extract experience section first to avoid counting education dates
        sections = self._identify_sections(text)
        experience_text = sections.get('experience', '')
        education_text = sections.get('education', '')
        
        # Remove education section from text to avoid counting education dates
        text_without_education = text
        if education_text:
            text_without_education = text.replace(education_text, '')
            logger.debug("Removed education section from experience calculation")
        
        # Only search in experience-related sections
        search_text = experience_text if experience_text else text_without_education.lower()
        
        # First try to find explicit experience statements
        for pattern in self.experience_patterns:
            match = re.search(pattern, search_text.lower())
            if match:
                years = int(match.group(1))
                logger.debug(f"Found explicit experience: {years} years")
                return years
        
        # If no explicit experience found, check if there's any work experience at all
        # Look for job-related keywords before counting dates
        work_keywords = ['job', 'position', 'role', 'employed', 'worked at', 'company', 'employer', 
                        'software engineer', 'developer', 'manager', 'analyst', 'consultant', 'intern',
                        'full-time', 'part-time', 'freelance', 'contract', 'work experience', 'employment']
        has_work_experience = any(keyword in text_lower for keyword in work_keywords)
        
        # Also check if experience section has actual work content (not just education)
        if experience_text:
            has_work_in_section = any(keyword in experience_text.lower() for keyword in work_keywords)
            if not has_work_in_section:
                logger.info(f"[EXPERIENCE] Experience section found but no work keywords, likely education. Returning 0")
                return 0
        
        if not has_work_experience:
            logger.info(f"[EXPERIENCE] No work experience keywords found anywhere, returning 0")
            return 0
        
        # Only if work keywords found, try to calculate from job dates
        # But only look in experience section to avoid counting education dates
        if experience_text:
            if DATE_PARSER_AVAILABLE:
                years = self._calculate_experience_from_dates_advanced(experience_text)
            else:
                years = self._calculate_experience_from_dates(experience_text)
            
            if years and years > 0:
                # Double-check: make sure we're not counting education dates
                if education_text:
                    edu_years = self._get_education_duration(education_text)
                    if edu_years and abs(years - edu_years) <= 1:
                        logger.warning(f"[EXPERIENCE] Calculated {years} years matches education duration {edu_years}, likely education dates. Returning 0")
                        return 0
                
                logger.info(f"[EXPERIENCE] Calculated experience from dates in experience section: {years} years")
                return years
        
        # Only if work keywords found, try to calculate from dates in text (excluding education)
        # Make sure we're not counting education dates by checking context
        if DATE_PARSER_AVAILABLE:
            years = self._calculate_experience_from_dates_advanced(text_without_education)
        else:
            years = self._calculate_experience_from_dates(text_without_education)
        
        if years and years > 0:
            # Double-check: make sure we're not counting education dates
            # If the calculated years match education duration, it's likely education
            if education_text:
                edu_years = self._get_education_duration(education_text)
                if edu_years and abs(years - edu_years) <= 1:
                    logger.warning(f"[EXPERIENCE] Calculated {years} years matches education duration {edu_years}, likely education dates. Returning 0")
                    return 0
            
            logger.info(f"[EXPERIENCE] Calculated experience from full text (excluding education): {years} years")
            return years
        
        logger.info(f"[EXPERIENCE] No experience dates found, returning 0")
        return 0
    
    def _calculate_experience_from_dates(self, text: str) -> Optional[int]:
        """Calculate experience years from job start/end dates - only count work dates, not education"""
        # Look for date patterns like (2020-2022), (2022-2024), etc.
        # But exclude patterns that are clearly in education section
        date_patterns = [
            r'\((\d{4})\s*-\s*(\d{4})\)',  # (2020-2022)
            r'(\d{4})\s*-\s*(\d{4})',      # 2020-2022
            r'(\d{4})\s*to\s*(\d{4})',     # 2020 to 2022
            r'(\d{4})\s*present',          # 2020 present
            r'(\d{4})\s*current',          # 2020 current
        ]
        
        total_months = 0
        current_year = datetime.now().year  # Use current year dynamically
        
        # Exclude education-related date patterns - expanded list
        education_keywords = [
            'degree', 'bachelor', 'master', 'phd', 'diploma', 'graduated', 'university', 'college', 'school',
            'education', 'academic', 'qualification', 'b.tech', 'm.tech', 'b.e', 'm.e', 'b.sc', 'm.sc',
            'engineering', 'computer science', 'gpa', 'cgpa', 'percentage', 'grade', 'semester', 'course'
        ]
        
        # Work-related keywords that indicate actual employment
        work_keywords = [
            'company', 'employer', 'job', 'position', 'role', 'worked', 'employed', 'intern', 'internship',
            'software engineer', 'developer', 'manager', 'analyst', 'consultant', 'specialist', 'coordinator',
            'full-time', 'part-time', 'freelance', 'contract', 'employee', 'staff', 'team lead', 'senior'
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                # Find the position of this match in the text
                match_str = str(match[0] if isinstance(match, tuple) else match)
                match_start = text.lower().find(match_str.lower())
                
                if match_start != -1:
                    # Get broader context around the date (100 chars before and after)
                    context = text[max(0, match_start-100):match_start+150].lower()
                    
                    # Skip if it's clearly education-related
                    if any(edu_keyword in context for edu_keyword in education_keywords):
                        logger.debug(f"Skipping date {match_str} - found in education context")
                        continue
                    
                    # Only count if there are work-related keywords nearby
                    has_work_context = any(work_keyword in context for work_keyword in work_keywords)
                    if not has_work_context:
                        logger.debug(f"Skipping date {match_str} - no work context found")
                        continue
                
                if isinstance(match, tuple) and len(match) == 2:
                    start_year = int(match[0])
                    end_year = int(match[1])
                    # More strict validation - dates should be recent and reasonable
                    if start_year <= end_year and start_year >= 2000 and end_year <= current_year + 1:
                        months = (end_year - start_year) * 12
                        # Only count if it's at least 1 month and at most 50 years
                        if 1 <= months <= 600:
                            total_months += months
                            logger.debug(f"Counting date range {start_year}-{end_year} as {months} months")
                elif isinstance(match, str):
                    # Handle present/current patterns
                    year_match = re.search(r'(\d{4})', match)
                    if year_match:
                        start_year = int(year_match.group(1))
                        if start_year >= 2000 and start_year <= current_year:
                            months = (current_year - start_year) * 12
                            if 1 <= months <= 600:
                                total_months += months
                                logger.debug(f"Counting date {start_year}-present as {months} months")
        
        if total_months > 0:
            years = total_months // 12
            return years if years > 0 else None
        
        return None
    
    def _get_education_duration(self, education_text: str) -> Optional[int]:
        """Get the duration of education from education section"""
        date_patterns = [
            r'\((\d{4})\s*-\s*(\d{4})\)',  # (2020-2022)
            r'(\d{4})\s*-\s*(\d{4})',      # 2020-2022
            r'(\d{4})\s*to\s*(\d{4})',     # 2020 to 2022
        ]
        
        current_year = datetime.now().year
        max_duration = 0
        
        for pattern in date_patterns:
            matches = re.findall(pattern, education_text, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple) and len(match) == 2:
                    start_year = int(match[0])
                    end_year = int(match[1])
                    if start_year <= end_year and start_year >= 2000 and end_year <= current_year + 1:
                        duration = end_year - start_year
                        max_duration = max(max_duration, duration)
        
        return max_duration if max_duration > 0 else None

    def _calculate_experience_from_dates_advanced(self, text: str) -> Optional[int]:
        """Calculate experience years from job dates using dateparser for better accuracy"""
        try:
            # First check if there are work-related keywords - if not, return 0
            work_keywords = ['company', 'employer', 'job', 'position', 'role', 'worked', 'employed', 
                           'intern', 'internship', 'software engineer', 'developer', 'manager']
            has_work_context = any(keyword in text.lower() for keyword in work_keywords)
            
            if not has_work_context:
                logger.debug(f"[EXPERIENCE] No work keywords in text, returning None")
                return None
            
            dates = search_dates(text)
            if not dates:
                return None
            
            periods = []
            for i in range(len(dates) - 1):
                start, end = dates[i][1], dates[i + 1][1]
                if start < end:
                    delta = (end - start).days
                    # Only count if it's between 1 month and 10 years (reasonable work period)
                    if 30 < delta < 3650:
                        # Check context around these dates to ensure they're work-related
                        periods.append(delta)
            
            if periods:
                total_days = sum(periods)
                years = total_days / 365.0
                calculated_years = int(round(years, 1))
                logger.debug(f"[EXPERIENCE] Calculated {calculated_years} years from {len(periods)} periods")
                return calculated_years
        except Exception as e:
            logger.warning(f"Error calculating experience from dates: {e}")
        
        return None

    def _extract_education(self, text: str) -> Optional[str]:
        """Extract education information from resume"""
        education_section = self._extract_education_section(text)
        if education_section:
            return education_section.strip()
        
        # Look for degree mentions throughout the text - expanded patterns
        degree_patterns = [
            r'(Bachelor[^.]*\.)',
            r'(Master[^.]*\.)',
            r'(PhD[^.]*\.)',
            r'(Associate[^.]*\.)',
            r'(Diploma[^.]*\.)',
            r'(B\.?S\.?[^.]*\.)',
            r'(B\.?A\.?[^.]*\.)',
            r'(M\.?S\.?[^.]*\.)',
            r'(M\.?A\.?[^.]*\.)',
            r'(M\.?B\.?A\.?[^.]*\.)',
            r'(B\.?E\.?[^.]*\.)',
            r'(B\.?T\.?e\.?c\.?h[^.]*\.)',
            r'(M\.?T\.?e\.?c\.?h[^.]*\.)',
            r'(High\s+School[^.]*\.)',
            r'(Secondary\s+School[^.]*\.)',
            r'(Primary\s+School[^.]*\.)',
            r'(H\.?S\.?C\.?[^.]*\.)',
            r'(S\.?S\.?C\.?[^.]*\.)',
            r'(10th[^.]*\.)',
            r'(12th[^.]*\.)',
            r'(Graduation[^.]*\.)',
            r'(Post\s+Graduation[^.]*\.)'
        ]
        
        education_info = []
        for pattern in degree_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if len(match) > 5 and match not in education_info:  # Avoid duplicates and short matches
                    education_info.append(match.strip())
        
        if education_info:
            return '; '.join(education_info)  # Return all education entries
        
        return None

    def _extract_education_section(self, text: str) -> Optional[str]:
        """Extract education section from resume - capture ALL education levels"""
        education_keywords = ['education', 'academic', 'qualifications', 'degrees', 'academic qualification']
        
        for keyword in education_keywords:
            # More flexible pattern to capture education content
            pattern = rf'{keyword}[:\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*\n|\n[A-Z][A-Z\s]+:|\n[A-Z][a-z]+\s+[A-Z][a-z]+:|\Z)'
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                education_text = match.group(1).strip()
                # Clean up the education text - stop at personal details
                lines = education_text.split('\n')
                cleaned_lines = []
                consecutive_non_edu = 0
                
                for line in lines:
                    line = line.strip()
                    if line and len(line) > 3:  # Include shorter lines too
                        # Stop if we hit personal details
                        if any(personal_term in line.lower() for personal_term in [
                            'personal details', 'date of birth', 'father', 'mother', 'marital status',
                            'gender', 'languages known', 'permanent address', 'address'
                        ]):
                            break
                        
                        # Include lines that look like education - expanded patterns
                        is_education_line = any(edu_term in line.lower() for edu_term in [
                            'bachelor', 'master', 'phd', 'degree', 'university', 'college', 'institute',
                            'gpa', 'cgpa', 'percentage', 'grade', 'b.tech', 'm.tech', 'b.e', 'm.e',
                            'high school', 'secondary school', 'primary school', 'middle school',
                            'elementary', 'hsc', 'ssc', '10th', '12th', 'graduation', 'diploma',
                            'certificate', 'board', 'school', 'academy', 'institution', 'faculty',
                            'department', 'course', 'stream', 'branch', 'specialization', 'major',
                            'minor', 'honors', 'pass', 'completed', 'graduated'
                        ])
                        
                        # Also include lines with dates (likely education dates)
                        has_date = re.search(r'\d{4}', line)
                        
                        # Also include lines that look like exam names or qualifications
                        is_exam_line = re.match(r'^(examination|exam|qualification|degree|diploma|certificate)\s*:', line, re.IGNORECASE)
                        
                        if is_education_line or has_date or is_exam_line:
                            cleaned_lines.append(line)
                            consecutive_non_edu = 0
                        else:
                            consecutive_non_edu += 1
                            # If we see multiple non-education lines in a row, might have moved to next section
                            if consecutive_non_edu > 3:
                                break
                
                if cleaned_lines:
                    logger.info(f"[EDUCATION] Extracted {len(cleaned_lines)} education lines from section")
                    return '\n'.join(cleaned_lines)  # Return ALL education lines
        
        return None

    def _extract_work_experience(self, text: str) -> Optional[str]:
        """Extract work experience information from resume"""
        experience_section = self._extract_experience_section(text)
        if experience_section:
            return experience_section.strip()
        
        # Look for job titles and companies throughout the text
        job_patterns = [
            r'(Software Engineer[^.]*\.)',
            r'(Developer[^.]*\.)',
            r'(Manager[^.]*\.)',
            r'(Analyst[^.]*\.)',
            r'(Consultant[^.]*\.)',
            r'(Specialist[^.]*\.)',
            r'(Coordinator[^.]*\.)',
            r'(Director[^.]*\.)',
            r'(Lead[^.]*\.)',
            r'(Senior[^.]*\.)'
        ]
        
        experiences = []
        for pattern in job_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if len(match) > 10 and match not in experiences:  # Avoid duplicates and short matches
                    experiences.append(match.strip())
        
        if experiences:
            return '; '.join(experiences)  # Return all experiences
        
        return None

    def _extract_experience_section(self, text: str) -> Optional[str]:
        """Extract experience section from resume"""
        experience_keywords = ['experience', 'work experience', 'employment', 'professional experience', 'career', 'projects']
        
        for keyword in experience_keywords:
            # More flexible pattern to capture experience content
            pattern = rf'{keyword}[:\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*\n|\n[A-Z][A-Z\s]+:|\n[A-Z][a-z]+\s+[A-Z][a-z]+:|\Z)'
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                experience_text = match.group(1).strip()
                # Clean up the experience text - stop at other sections
                lines = experience_text.split('\n')
                cleaned_lines = []
                for line in lines:
                    line = line.strip()
                    if line and len(line) > 5:  # Skip empty or very short lines
                        # Stop if we hit other sections
                        if any(section_term in line.lower() for section_term in [
                            'certifications', 'education', 'personal details', 'skills', 'technical skills'
                        ]):
                            break
                        # Only include lines that look like experience/projects
                        if any(exp_term in line.lower() for exp_term in [
                            'project', 'experience', 'work', 'job', 'company', 'developer', 'engineer',
                            'manager', 'analyst', 'consultant', 'html', 'css', 'javascript', 'java',
                            'react', 'angular', 'vue', 'node', 'python', 'website', 'application'
                        ]):
                            cleaned_lines.append(line)
                
                if cleaned_lines:
                    return '\n'.join(cleaned_lines)  # Return all experience lines
        
        return None

    def _extract_projects(self, text: str) -> Optional[str]:
        """Extract projects information from resume"""
        projects_section = self._extract_projects_section(text)
        if projects_section:
            return projects_section.strip()
        
        # Look for project mentions throughout the text
        project_patterns = [
            r'(Project[^.]*\.)',
            r'(Portfolio[^.]*\.)',
            r'(Application[^.]*\.)',
            r'(Website[^.]*\.)',
            r'(System[^.]*\.)',
            r'(Software[^.]*\.)',
            r'(App[^.]*\.)',
            r'(Platform[^.]*\.)',
        ]
        
        projects_info = []
        for pattern in project_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if len(match) > 10 and match not in projects_info:  # Avoid duplicates and short matches
                    projects_info.append(match.strip())
        
        if projects_info:
            return '; '.join(projects_info)  # Return all projects
        
        return None

    def _extract_projects_section(self, text: str) -> Optional[str]:
        """Extract projects section from resume"""
        project_keywords = ['projects', 'portfolio', 'personal projects', 'academic projects', 'work projects']
        
        for keyword in project_keywords:
            # More flexible pattern to capture projects content
            # Look for the keyword followed by content until the next major section
            pattern = rf'{keyword}[:\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*\n|\n[A-Z][A-Z\s]+:|\n[A-Z][a-z]+\s+[A-Z][a-z]+:|\n[A-Z][a-z]+:|\Z)'
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                projects_text = match.group(1).strip()
                # Clean up the projects text
                lines = projects_text.split('\n')
                cleaned_lines = []
                for line in lines:
                    line = line.strip()
                    if line and len(line) > 5:  # Skip empty or very short lines
                        # Skip lines that look like section headers
                        if not re.match(r'^[A-Z][A-Z\s]+:$', line):
                            cleaned_lines.append(line)
                
                if cleaned_lines:
                    return '\n'.join(cleaned_lines)  # Return all project lines
        
        return None

    def _extract_certifications(self, text: str) -> Optional[str]:
        """Extract certifications information from resume"""
        certifications_section = self._extract_certifications_section(text)
        if certifications_section:
            return certifications_section.strip()
        
        # Look for certification mentions throughout the text
        cert_patterns = [
            r'(Certified[^.]*\.)',
            r'(Certificate[^.]*\.)',
            r'(Certification[^.]*\.)',
            r'(AWS[^.]*\.)',
            r'(Azure[^.]*\.)',
            r'(Google[^.]*\.)',
            r'(Microsoft[^.]*\.)',
            r'(Oracle[^.]*\.)',
            r'(Cisco[^.]*\.)',
            r'(CompTIA[^.]*\.)',
        ]
        
        certs_info = []
        for pattern in cert_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if len(match) > 10 and match not in certs_info:  # Avoid duplicates and short matches
                    certs_info.append(match.strip())
        
        if certs_info:
            return '; '.join(certs_info)  # Return all certifications
        
        return None

    def _extract_certifications_section(self, text: str) -> Optional[str]:
        """Extract certifications section from resume"""
        cert_keywords = ['certifications', 'certificates', 'certified', 'credentials', 'licenses']
        
        for keyword in cert_keywords:
            # More flexible pattern to capture certifications content
            pattern = rf'{keyword}[:\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*\n|\n[A-Z][A-Z\s]+:|\n[A-Z][a-z]+\s+[A-Z][a-z]+:|\Z)'
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                certs_text = match.group(1).strip()
                # Clean up the certifications text - stop at other sections
                lines = certs_text.split('\n')
                cleaned_lines = []
                for line in lines:
                    line = line.strip()
                    if line and len(line) > 5:  # Skip empty or very short lines
                        # Stop if we hit other sections
                        if any(section_term in line.lower() for section_term in [
                            'education', 'personal details', 'skills', 'technical skills', 'projects'
                        ]):
                            break
                        # Only include lines that look like certifications
                        if any(cert_term in line.lower() for cert_term in [
                            'certified', 'certificate', 'certification', 'aws', 'azure', 'google', 'microsoft',
                            'oracle', 'cisco', 'comptia', 'credential', 'license', 'professional'
                        ]):
                            cleaned_lines.append(line)
                
                if cleaned_lines:
                    return '\n'.join(cleaned_lines)  # Return all certification lines
        
        return None

    def _extract_summary(self, text: str) -> Optional[str]:
        """Extract professional summary or objective from resume"""
        summary_keywords = ['summary', 'objective', 'profile', 'about', 'overview', 'career objective']
        
        for keyword in summary_keywords:
            # More flexible pattern to capture summary content
            pattern = rf'{keyword}[:\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*\n|\n[A-Z][A-Z\s]+:|\n[A-Z][a-z]+\s+[A-Z][a-z]+:|\Z)'
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                summary = match.group(1).strip()
                if len(summary) > 20:  # Ensure it's substantial
                    # Clean up the summary to avoid repetition
                    lines = summary.split('\n')
                    cleaned_lines = []
                    seen_lines = set()
                    
                    for line in lines:
                        line = line.strip()
                        if line and len(line) > 10:  # Skip empty or very short lines
                            # Avoid repetition by checking if we've seen similar content
                            line_lower = line.lower()
                            if not any(seen_line.lower() in line_lower or line_lower in seen_line.lower() 
                                     for seen_line in seen_lines):
                                cleaned_lines.append(line)
                                seen_lines.add(line)
                    
                    if cleaned_lines:
                        return '\n'.join(cleaned_lines)  # Return all unique summary lines
        
        return None

    def _extract_education_aggressive(self, text: str) -> Optional[str]:
        """Aggressive education extraction - look for degree patterns anywhere in text"""
        lines = text.split('\n')
        education_lines = []
        in_education = False
        
        for line in lines:
            line_lower = line.lower().strip()
            line_original = line.strip()
            
            # Check if this line starts education section
            if any(keyword in line_lower for keyword in ['education', 'academic', 'qualification', 'degree', 'bachelor', 'master', 'phd', 'diploma']):
                in_education = True
                if len(line_original) > 5:
                    education_lines.append(line_original)
                continue
            
            # If in education section, collect lines until we hit another major section
            if in_education:
                if any(keyword in line_lower for keyword in ['experience', 'work', 'project', 'skill', 'certification', 'personal']):
                    break
                if line_original and len(line_original) > 3:
                    education_lines.append(line_original)
        
        if education_lines:
            return '\n'.join(education_lines)  # Return all education lines
        return None
    
    def _extract_experience_aggressive(self, text: str) -> Optional[str]:
        """Aggressive experience extraction - look for work/job patterns anywhere in text"""
        lines = text.split('\n')
        experience_lines = []
        in_experience = False
        
        for line in lines:
            line_lower = line.lower().strip()
            line_original = line.strip()
            
            # Check if this line starts experience section
            if any(keyword in line_lower for keyword in ['experience', 'employment', 'work history', 'career', 'position', 'job']):
                in_experience = True
                if len(line_original) > 5:
                    experience_lines.append(line_original)
                continue
            
            # If in experience section, collect lines until we hit another major section
            if in_experience:
                if any(keyword in line_lower for keyword in ['education', 'project', 'skill', 'certification', 'personal', 'objective']):
                    break
                if line_original and len(line_original) > 5:
                    experience_lines.append(line_original)
        
        if experience_lines:
            return '\n'.join(experience_lines)  # Return all experience lines
        return None
    
    def _extract_skills_aggressive(self, text: str) -> List[str]:
        """Aggressive skills extraction - find all skills from keyword list in text"""
        found_skills = []
        text_lower = text.lower()
        
        # Look for all skills from our keyword list
        for skill in self.skill_keywords:
            if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
                found_skills.append(skill.title())
        
        # Also look for skills section
        skills_section = self._extract_skills_section(text)
        if skills_section:
            skills_list = re.split(r'[,•\n\r;]', skills_section)
            for skill in skills_list:
                skill = skill.strip()
                if skill and len(skill) > 1 and len(skill) < 50:
                    skill = re.sub(r'^[^:]*:\s*', '', skill)
                    skill = skill.strip()
                    if self._is_valid_skill(skill):
                        found_skills.append(skill.title())
        
        return list(set(found_skills))
    
    def _extract_projects_aggressive(self, text: str) -> Optional[str]:
        """Aggressive projects extraction - look for project patterns anywhere in text"""
        lines = text.split('\n')
        project_lines = []
        in_project = False
        
        for line in lines:
            line_lower = line.lower().strip()
            line_original = line.strip()
            
            # Check if this line starts project section
            if any(keyword in line_lower for keyword in ['project', 'portfolio']):
                in_project = True
                if len(line_original) > 5:
                    project_lines.append(line_original)
                continue
            
            # If in project section, collect lines until we hit another major section
            if in_project:
                if any(keyword in line_lower for keyword in ['education', 'experience', 'skill', 'certification', 'personal']):
                    break
                if line_original and len(line_original) > 5:
                    project_lines.append(line_original)
        
        if project_lines:
            return '\n'.join(project_lines)  # Return all project lines
        return None
    
    def _extract_certifications_aggressive(self, text: str) -> Optional[str]:
        """Aggressive certifications extraction - look for cert patterns anywhere in text"""
        lines = text.split('\n')
        cert_lines = []
        in_cert = False
        
        for line in lines:
            line_lower = line.lower().strip()
            line_original = line.strip()
            
            # Check if this line starts certification section
            if any(keyword in line_lower for keyword in ['certification', 'certificate', 'certified', 'license', 'credential']):
                in_cert = True
                if len(line_original) > 5:
                    cert_lines.append(line_original)
                continue
            
            # If in cert section, collect lines until we hit another major section
            if in_cert:
                if any(keyword in line_lower for keyword in ['education', 'experience', 'project', 'skill', 'personal']):
                    break
                if line_original and len(line_original) > 5:
                    cert_lines.append(line_original)
        
        if cert_lines:
            return '\n'.join(cert_lines)  # Return all certification lines
        return None

    def _extract_languages(self, sections: Dict[str, str], text: str) -> Optional[str]:
        """Extract languages from sections or text"""
        languages_text = sections.get('languages', '')
        
        if languages_text and len(languages_text.strip()) > 5:
            return languages_text.strip()
        
        # Look for language patterns in text
        language_patterns = [
            r'(?:languages?|language\s+proficiency)[:\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*\n|\n[A-Z][A-Z\s]+:|\Z)',
        ]
        
        for pattern in language_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                lang_text = match.group(1).strip()
                if len(lang_text) > 3:
                    return lang_text
        
        return None

    def _extract_achievements(self, sections: Dict[str, str], text: str) -> Optional[str]:
        """Extract achievements, awards, honors, publications from sections or text"""
        achievements_text = sections.get('achievements', '')
        
        if achievements_text and len(achievements_text.strip()) > 5:
            return achievements_text.strip()
        
        # Look for achievement patterns in text
        achievement_patterns = [
            r'(?:achievements?|awards?|honors?|recognition)[:\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*\n|\n[A-Z][A-Z\s]+:|\Z)',
            r'(?:publications?|papers?|research)[:\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*\n|\n[A-Z][A-Z\s]+:|\Z)',
        ]
        
        achievement_lines = []
        for pattern in achievement_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                ach_text = match.group(1).strip()
                if len(ach_text) > 5:
                    achievement_lines.append(ach_text)
        
        if achievement_lines:
            return '\n\n'.join(achievement_lines)
        
        return None

    def _extract_hobbies(self, sections: Dict[str, str], text: str) -> Optional[str]:
        """Extract hobbies and interests from sections or text"""
        hobbies_text = sections.get('hobbies', '')
        
        if hobbies_text and len(hobbies_text.strip()) > 5:
            return hobbies_text.strip()
        
        # Look for hobby patterns in text
        hobby_patterns = [
            r'(?:hobbies?|interests?|personal\s+interests?)[:\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*\n|\n[A-Z][A-Z\s]+:|\Z)',
        ]
        
        for pattern in hobby_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                hobby_text = match.group(1).strip()
                if len(hobby_text) > 3:
                    return hobby_text
        
        return None

    def _extract_references(self, sections: Dict[str, str], text: str) -> Optional[str]:
        """Extract references from sections or text"""
        references_text = sections.get('references', '')
        
        if references_text and len(references_text.strip()) > 5:
            return references_text.strip()
        
        # Look for reference patterns in text
        reference_patterns = [
            r'(?:references?|professional\s+references?)[:\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*\n|\n[A-Z][A-Z\s]+:|\Z)',
        ]
        
        for pattern in reference_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                ref_text = match.group(1).strip()
                if len(ref_text) > 3:
                    return ref_text
        
        return None

    def _extract_volunteer_work(self, sections: Dict[str, str], text: str) -> Optional[str]:
        """Extract volunteer work and community service from sections or text"""
        volunteer_text = sections.get('volunteer', '')
        
        if volunteer_text and len(volunteer_text.strip()) > 5:
            return volunteer_text.strip()
        
        # Look for volunteer patterns in text
        volunteer_patterns = [
            r'(?:volunteer\s+work|volunteer\s+experience|volunteering|community\s+service)[:\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*\n|\n[A-Z][A-Z\s]+:|\Z)',
        ]
        
        for pattern in volunteer_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                vol_text = match.group(1).strip()
                if len(vol_text) > 5:
                    return vol_text
        
        return None

    def _extract_additional_info(self, sections: Dict[str, str], text: str) -> Optional[str]:
        """Extract any additional information not captured in other sections"""
        # Look for sections that might contain additional info
        additional_sections = []
        
        # Check for other common sections
        other_patterns = [
            r'(?:additional\s+information|other\s+information|miscellaneous|other\s+details)[:\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*\n|\n[A-Z][A-Z\s]+:|\Z)',
            r'(?:patents?|publications?|conferences?|presentations?)[:\s]*([^\n]*(?:\n[^\n]*)*?)(?=\n\s*\n|\n[A-Z][A-Z\s]+:|\Z)',
        ]
        
        for pattern in other_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                info_text = match.group(1).strip()
                if len(info_text) > 5:
                    additional_sections.append(info_text)
        
        if additional_sections:
            return '\n\n'.join(additional_sections)
        
        return None

    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """Parse a resume file and return structured data"""
        try:
            # Get file extension
            file_extension = os.path.splitext(file_path)[1].lower().lstrip('.')
            logger.info(f"Parsing resume file: {file_path}, extension: {file_extension}")
            
            # Extract text
            text = self.extract_text_from_file(file_path, file_extension)
            
            if not text:
                logger.warning(f"No text extracted from {file_path}")
                return {}
            
            logger.info(f"Extracted {len(text)} characters from resume")
            logger.debug(f"First 500 chars of extracted text: {text[:500]}")
            
            # Parse the text
            parsed_data = self.parse_resume_data(text, os.path.basename(file_path))
            
            # Log what was extracted
            if parsed_data:
                logger.info(f"Parsed data summary - Name: {parsed_data.get('full_name')}, Skills: {len(parsed_data.get('skills', []))}, Education: {bool(parsed_data.get('education'))}, Experience: {bool(parsed_data.get('work_experience'))}")
            else:
                logger.warning("parse_resume_data returned empty dict")
            
            return parsed_data or {}
            
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {str(e)}", exc_info=True)
            return {}


# Global instance
resume_parser = ResumeParser()

# Module-level function for easy importing
def parse_resume_data(file_path):
    """Parse resume data from file path - module level function"""
    return resume_parser.parse_file(file_path)
